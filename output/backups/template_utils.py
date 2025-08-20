from __future__ import annotations
import re
import unicodedata
import logging
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional

# optional deps
try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

try:
    import jinja2
except ModuleNotFoundError:
    jinja2 = None

# enable debug logging here for template utils (useful while developing)
logging.basicConfig(level=logging.DEBUG, format="%(levelname)s: %(message)s")

# Email regex used by find_email_in_record
EMAIL_REGEX = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", re.UNICODE)


def _clean_cell_value(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, float):
        if val.is_integer():
            s = str(int(val))
        else:
            s = format(val, "f").rstrip("0").rstrip(".")
    else:
        s = str(val)
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\u3000", " ")
    s = re.sub(r"[\u200B-\u200F\uFEFF]", "", s)
    s = re.sub(r"[\x00-\x1F\x7F]", "", s)
    return s.strip()


def find_email_in_record(record: Dict[str, Any]) -> Optional[str]:
    """
    Robustly find an email address in a record's fields.
    """
    # 1) scan keys that look mail-like
    for k in list(record.keys()):
        if k is None:
            continue
        k_norm = _clean_cell_value(k).lower()
        try:
            if re.search(r"(mail|email|e-?mail|信箱|電子郵)", k_norm):
                raw = record.get(k)
                s = _clean_cell_value(raw)
                if not s:
                    continue
                parts = re.split(r"[;,/|\s]+", s)
                for p in parts:
                    p = p.strip().strip("()[]<>\"'")
                    if not p:
                        continue
                    if EMAIL_REGEX.fullmatch(p):
                        return p
                    m = EMAIL_REGEX.search(p)
                    if m:
                        return m.group(0)
                m = EMAIL_REGEX.search(s)
                if m:
                    return m.group(0)
        except Exception:
            continue

    # 2) fallback: scan all values
    for v in record.values():
        try:
            s = _clean_cell_value(v)
            if not s:
                continue
            parts = re.split(r"[;,/|\s]+", s)
            for p in parts:
                p = p.strip().strip("()[]<>\"'")
                if not p:
                    continue
                if EMAIL_REGEX.fullmatch(p):
                    return p
                m = EMAIL_REGEX.search(p)
                if m:
                    return m.group(0)
            m = EMAIL_REGEX.search(s)
            if m:
                return m.group(0)
        except Exception:
            continue

    return None


def sanitize_filename(s: str, max_len: int = 200) -> str:
    s = re.sub(r"[\\/:\*\?\"<>\|]+", "-", s or "")
    s = re.sub(r"\s+", " ", s).strip()
    return s[:max_len]


def find_template_file(template_filename: str, template_dir: Optional[Path] = None) -> Path:
    """
    Find template file under template_dir (or default templates/).
    """
    # attempt to locate TEMPLATE_DIR similar to other scripts
    if template_dir is None:
        BASE_DIR = Path(__file__).resolve().parents[2]
        template_dir = BASE_DIR / "templates"
    p = Path(template_dir) / template_filename
    if p.exists():
        return p
    matches = list(Path(template_dir).rglob(template_filename))
    if matches:
        return matches[0]
    raise FileNotFoundError(f"找不到模板：{template_filename}（已搜尋 {template_dir} 及子資料夾）")


def render_docx_template(template_path: Path, out_path: Path, mapping: Dict[str, Any], replacers: Optional[List[Tuple[str, str]]] = None) -> None:
    """
    Render a .docx template by performing paragraph-level replacements.
    replacers: list of (pattern, mapping_key). If None, will use mapping keys in the form {{key}}.
    """
    if Document is None:
        raise ModuleNotFoundError("python-docx is required for render_docx_template")
    doc = Document(str(template_path))

    def apply_text(text: str) -> str:
        if not text:
            return text
        out = text
        if replacers:
            for pat, key in replacers:
                out = out.replace(pat, str(mapping.get(key, "")))
        else:
            # generic: replace {{ key }} or {{key}} for mapping keys
            for k, v in mapping.items():
                out = out.replace("{{" + str(k) + "}}", str(v))
                out = out.replace("{{ " + str(k) + " }}", str(v))
        return out

    # paragraphs
    for para in doc.paragraphs:
        if "{{" in para.text or "}}" in para.text:
            new = apply_text(para.text)
            # 如果 template 裡有人輸入 literal "\n"，把它轉成真正的換行字元
            if new is not None:
                new = new.replace("\\n", "\n")
            if new != para.text:
                if para.runs:
                    para.runs[0].text = new
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(new)

    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text or "}}" in para.text:
                        new = apply_text(para.text)
                        if new is not None:
                            new = new.replace("\\n", "\n")
                        if new != para.text:
                            if para.runs:
                                para.runs[0].text = new
                                for r in para.runs[1:]:
                                    r.text = ""
                            else:
                                para.add_run(new)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def render_body_from_template(template_path: Path, context: Dict[str, Any]) -> str:
    """
    Render textual body from a docx template. Returns string (paragraphs joined with \n).
    Uses Jinja2 if available, else falls back to manual placeholder resolution.
    Handles run-splitting by rendering the whole paragraph text and writing back to the first run.
    Also converts literal backslash+n ("\\n") into real newline characters.
    """
    if Document is None:
        raise ModuleNotFoundError("python-docx is required for template rendering")
    doc = Document(str(template_path))

    jenv = None
    if jinja2 is not None:
        try:
            jenv = jinja2.Environment(undefined=jinja2.StrictUndefined)
        except Exception:
            jenv = None

    # flatten context into dotted keys for simple replacement
    flat: Dict[str, str] = {}

    def _flatten(prefix: str, val: Any) -> None:
        if isinstance(val, dict):
            for k, v in val.items():
                _flatten(f"{prefix}.{k}" if prefix else str(k), v)
        elif isinstance(val, (list, tuple)):
            for i, item in enumerate(val):
                _flatten(f"{prefix}[{i}]" if prefix else f"[{i}]", item)
        else:
            flat[prefix] = "" if val is None else str(val)

    for k, v in context.items():
        _flatten(str(k), v)

    placeholder_re = re.compile(r"\{\{\s*(.*?)\s*\}\}")

    def render_text(text: str) -> str:
        if not text:
            return text
        # try jinja2 first (paragraph-level)
        if jenv is not None:
            try:
                tmpl = jenv.from_string(text)
                return tmpl.render(**context)
            except Exception:
                # don't crash; fall back to simple replacement
                pass

        # manual per-placeholder replacement using flat
        def _replace(match: re.Match) -> str:
            key = re.sub(r"\s+", "", match.group(1))
            return flat.get(key, match.group(0))

        return placeholder_re.sub(_replace, text)

    # Replace per-paragraph (not per-run) to avoid run-splitting issues
    for para in doc.paragraphs:
        if "{{" in para.text or "{%" in para.text:
            try:
                new_text = render_text(para.text)
            except Exception:
                new_text = para.text
            # 把 literal "\n" 轉成真換行（保險）
            if new_text is not None:
                new_text = new_text.replace("\\n", "\n")
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(new_text)

    # tables / cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text or "{%" in para.text:
                        try:
                            new_text = render_text(para.text)
                        except Exception:
                            new_text = para.text
                        if new_text is not None:
                            new_text = new_text.replace("\\n", "\n")
                        if para.runs:
                            para.runs[0].text = new_text
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.add_run(new_text)

    body_lines = [p.text for p in doc.paragraphs]
    return "\n".join(body_lines).strip()
