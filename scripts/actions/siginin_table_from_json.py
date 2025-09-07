from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import sys
# 若要把單元格直向置中（通常與固定高度一起用）
from docx.enum.table import WD_ALIGN_VERTICAL
from __init__ import format_date


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
# Project helpers
from scripts.core.bootstrap import DATA_DIR, OUTPUT_DIR, initialize
from scripts.actions.influencer import build_people
from docx.shared import Cm
from docx.shared import Pt, RGBColor

# -----------------------
# 可修改的程式碼參數（直接在程式中改）
TITLE_PT = 16                 # 全域字型大小 (pt)
FONT_PT = 14                 # 全域字型大小 (pt)
LEFT_RIGHT_MARGIN_CM = 1.5   # 左右邊界 (cm)
COL0_FIXED_CM = 8.0         # 第一欄 (Topic) 固定寬度 (cm)
COL1_FIXED_CM = 8.0          # 第二欄 (Speaker) 固定寬度 (cm)
# 第三欄由 available - (col0+col1) 決定
HDR_HEIGHT_CM = 0.9          # 表頭列高 (cm)
DATA_ROW_HEIGHT_CM = 3.5    # 每筆資料列高 (cm)
# DATA_ROWS_PER_PAGE = 20     # 若要每頁固定列數，可在 future 使用
# -----------------------

def load_program(program_id: int | None) -> Dict[str, Any]:
    data_file = DATA_DIR / "shared" / "program_data.json"
    programs_raw = json.loads(data_file.read_text(encoding="utf-8"))
    if isinstance(programs_raw, list):
        if program_id is not None:
            for prog in programs_raw:
                try:
                    if int(prog.get("id", -1)) == program_id:
                        return prog
                except (TypeError, ValueError):
                    continue
        return programs_raw[0] if programs_raw else {}
    elif isinstance(programs_raw, dict):
        return programs_raw
    return {}

def set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    run.font.size = Pt(size_pt)
    run.bold = bold

def set_table_cell_margins(table, left_cm: float = 0.1, right_cm: float = 0.1, top_pt: int = 0, bottom_pt: int = 0):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tcMar = tblPr.find(qn("w:tblCellMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tblCellMar")
        tblPr.append(tcMar)
    def _set_node(name: str, cm_val: float, parent: OxmlElement):
        node = parent.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            parent.append(node)
        node.set(qn("w:w"), str(int(cm_val * 567)))  # dxa
        node.set(qn("w:type"), "dxa")
    _set_node("left", left_cm, tcMar)
    _set_node("right", right_cm, tcMar)
    _set_node("top", (top_pt / 72.0) * 2.54 if top_pt else 0, tcMar)
    _set_node("bottom", (bottom_pt / 72.0) * 2.54 if bottom_pt else 0, tcMar)

def _get_first_nonempty(sp: dict, keys: List[str]) -> str:
    for k in keys:
        val = sp.get(k)
        if isinstance(val, str) and val.strip():
            return val.strip()
        if isinstance(val, dict):
            for subk in ("organization", "company", "affiliation", "dept", "department", "unit"):
                v2 = val.get(subk)
                if isinstance(v2, str) and v2.strip():
                    return v2.strip()
            for subk in ("title", "position", "role"):
                v2 = val.get(subk)
                if isinstance(v2, str) and v2.strip():
                    return v2.strip()
    return ""

def _set_table_total_width(table, total_cm: float):
    tbl = table._tbl
    tblPr = tbl.tblPr
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(total_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

def safe_set_row_height(row, height_cm: float, preferred_rule: str = "EXACT"):
    """設定 row.height，並安全嘗試設定 height_rule（各版本相容）。"""
    row.height = Cm(height_cm)
    rule_value = None
    if hasattr(WD_ROW_HEIGHT_RULE, preferred_rule):
        rule_value = getattr(WD_ROW_HEIGHT_RULE, preferred_rule)
    else:
        for alt in ("EXACT", "AT_LEAST", "AUTO"):
            if hasattr(WD_ROW_HEIGHT_RULE, alt):
                rule_value = getattr(WD_ROW_HEIGHT_RULE, alt)
                break
    if rule_value is not None:
        try:
            row.height_rule = rule_value
        except Exception:
            pass



def set_row_height_exact(row, height_cm: float):
    """
    對單一 row 設定固定高度 (cm)。如果 environment 支援 EXACT，使用 EXACT，
    否則退回 AT_LEAST / AUTO（安全相容）。
    """
    # 設高度（cm）
    row.height = Cm(height_cm)

    # 嘗試使用 EXACT，沒有再退回 AT_LEAST / AUTO
    rule_value = None
    if hasattr(WD_ROW_HEIGHT_RULE, "EXACT"):
        rule_value = WD_ROW_HEIGHT_RULE.EXACT
    elif hasattr(WD_ROW_HEIGHT_RULE, "AT_LEAST"):
        rule_value = WD_ROW_HEIGHT_RULE.AT_LEAST
    elif hasattr(WD_ROW_HEIGHT_RULE, "AUTO"):
        rule_value = WD_ROW_HEIGHT_RULE.AUTO

    if rule_value is not None:
        try:
            row.height_rule = rule_value
        except Exception:
            # 某些 python-docx 版本可能不允許直接指派，忽略錯誤即可
            pass

def set_table_rows_height(table, height_cm: float):
    """對 table 裡所有現存的 row 設定固定高度（包含 header 與資料列）"""
    for r in table.rows:
        set_row_height_exact(r, height_cm)


def set_cell_vertical_center(cell):
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        # 忽略不支援情況
        pass
def set_cell_background(cell, color_hex: str):
    """
    對單一 cell 設定背景色（hex，例如 '#BFBFBF' 或 'BFBFBF'）。
    """
    # normalize
    color = color_hex.lstrip("#")
    # access or create tcPr
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.append(tcPr)
    # remove existing shd if any
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    # create shading element
    shd = OxmlElement("w:shd")
    # w:val can be 'clear' or 'nil' etc; set fill to color hex
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color.upper())
    tcPr.append(shd)

def set_run_color_black(run):
    """簡單把 run 的文字顏色設成黑色，避免淺底看不清楚。"""
    try:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception:
        pass
def main() -> None:
    parser = argparse.ArgumentParser(description="Render program to docx")
    parser.add_argument("--program-id", type=int, default=None, help="Program id to render")
    parser.add_argument("--out", type=Path, default=None, help="Output .docx path")
    args = parser.parse_args()

    initialize()
    program = load_program(args.program_id)

    # load influencer & build people
    infl_file = DATA_DIR / "shared" / "influencer_data.json"
    try:
        influencers = json.loads(infl_file.read_text(encoding="utf-8"))
    except OSError:
        influencers = []
    chairs, speakers = build_people(program, influencers)

    # program topic map
    program_speakers = program.get("speakers", []) or []
    program_topic_map: Dict[str, str] = {
        (entry.get("name") or "").strip(): (entry.get("topic") or "").strip()
        for entry in program_speakers
    }

    planName = program.get("planName")
    event_name = (program.get("eventNames") or ["Program"])[0]
    date = program.get("date")

    out_path = args.out or (OUTPUT_DIR / f"講師簽到表_{program.get('eventNames[0]')}.docx")

    # 使用程式內參數 FONT_PT
    font_pt = int(FONT_PT)

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    normal_font.size = Pt(font_pt)

    # Title
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t_run = p_t.add_run(str(planName))
    set_run_font(p_t_run, TITLE_PT, bold=True)

    p_e = doc.add_paragraph()
    p_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_e_run = p_e.add_run(str(event_name))
    set_run_font(p_e_run, TITLE_PT, bold=True)

    p_e = doc.add_paragraph()
    slash_date =format_date(date, sep="/")
    p_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_e_run = p_e.add_run("("+str(slash_date)+")")
    set_run_font(p_e_run, TITLE_PT, bold=True)

    p_e = doc.add_paragraph()
    p_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_e_run = p_e.add_run("講員簽到單")
    set_run_font(p_e_run, TITLE_PT, bold=True)

    # 固定左右邊界（cm）
    doc.sections[0].left_margin = Cm(LEFT_RIGHT_MARGIN_CM)
    doc.sections[0].right_margin = Cm(LEFT_RIGHT_MARGIN_CM)

    # 建表（寫死欄寬但有保護機制）
    signin_table = doc.add_table(rows=1, cols=3, style="Table Grid")
    signin_table.autofit = False
    signin_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        signin_table.left_indent = Cm(0)
    except Exception:
        pass

    # 計算 page 與 available width（cm）
    page_width_cm = float(doc.sections[0].page_width) / float(Cm(1))
    available_cm = page_width_cm - (LEFT_RIGHT_MARGIN_CM * 2)

    # 先計算第三欄為剩餘空間
    col0 = float(COL0_FIXED_CM)
    col1 = float(COL1_FIXED_CM)
    col2 = available_cm - (col0 + col1)

    # 若總寬超過 available，等比例縮放三欄（確保不超出）
    total_req = col0 + col1 + (col2 if col2 > 0 else 0)
    if total_req > available_cm:
        scale = available_cm / total_req
        col0 *= scale
        col1 *= scale
        col2 = max(1.5, available_cm - (col0 + col1))

    # 最後保護（col2 至少 1.5 cm）
    if col2 < 1.5:
        col2 = 1.5
        if col0 + col1 + col2 > available_cm:
            remain = max(0, available_cm - col2)
            if col0 + col1 > 0:
                ratio = col0 / (col0 + col1)
            else:
                ratio = 0.6
            col0 = remain * ratio
            col1 = remain - col0

    # 設定欄寬與 table 寬
    signin_table.columns[0].width = Cm(col0)
    signin_table.columns[1].width = Cm(col1)
    signin_table.columns[2].width = Cm(col2)
    _set_table_total_width(signin_table, col0 + col1 + col2)

    # 縮小 cell padding
    set_table_cell_margins(signin_table, left_cm=0.12, right_cm=0.12)





    # 表頭
    hdr = signin_table.rows[0]
    safe_set_row_height(hdr, HDR_HEIGHT_CM)



    # 設 header 高度為 0.9 cm
    set_row_height_exact(signin_table.rows[0], 0.9)
    hdr_cells = hdr.cells
    p = hdr_cells[0].paragraphs[0]
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("主題 Topic")
    set_run_font(run, font_pt, bold=True)

    p = hdr_cells[1].paragraphs[0]
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("姓名 Name")
    set_run_font(run, font_pt, bold=True)

    p = hdr_cells[2].paragraphs[0]
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("簽到 Sign-in")
    set_run_font(run, font_pt, bold=True)


    for cell in hdr.cells:
        # set_table_cell_margins(cell, left_cm=0.0, right_cm=0.0, top_cm=0.0, bottom_cm=0.0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 並把 paragraph 設水平置中
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "#BFBFBF")
    # 資料列（每位講者一列）
    # 在檔案開頭確保已 import：
# from docx.enum.table import WD_ALIGN_VERTICAL

    # 替換原本的 if speakers: 迴圈
    if speakers:
        for sp in speakers:
            # 新增一列並設定固定列高（使用你已有的 helper）
            row = signin_table.add_row()
            try:
                set_row_height_exact(row, DATA_ROW_HEIGHT_CM)  # 若你定義了 exact helper
            except NameError:
                safe_set_row_height(row, DATA_ROW_HEIGHT_CM)

            row_cells = row.cells
            name_val = (sp.get("name") or "").strip()
            topic_val = program_topic_map.get(name_val, "") or ""

            # ---------- 第一欄：Topic（覆寫第一個 paragraph） ----------
            c0 = row_cells[0]
            p0 = c0.paragraphs[0] if c0.paragraphs else c0.add_paragraph()
            # 清除段內 runs（保險）
            for r in list(p0.runs):
                try:
                    r._element.getparent().remove(r._element)
                except Exception:
                    pass
            # 寫入 topic（若空放 NBSP）
            p0.add_run(topic_val or "\u00A0")
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 取消段落上下間距（可選）
            try:
                p0.paragraph_format.space_before = Pt(0)
                p0.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            # 垂直置中 cell
            try:
                c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass

            # ---------- 第二欄：Speaker（姓名+職稱 第一行，組織 第二行） ----------
            title_val = _get_first_nonempty(sp, ["title", "position", "role"]) or ""
            org_val = _get_first_nonempty(sp, ["organization", "company", "affiliation", "department", "dept", "unit", "employer"]) or ""

            c1 = row_cells[1]
            p1 = c1.paragraphs[0] if c1.paragraphs else c1.add_paragraph()
            # 清除原 runs
            for r in list(p1.runs):
                try:
                    r._element.getparent().remove(r._element)
                except Exception:
                    pass
            # 姓名（粗體）
            r_name = p1.add_run(name_val or "\u00A0")
            set_run_font(r_name, font_pt, bold=True)
            # 職稱（同段，不粗體）
            if title_val:
                r_title = p1.add_run(" " + title_val)
                set_run_font(r_title, font_pt, bold=False)
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p1.paragraph_format.space_before = Pt(0)
                p1.paragraph_format.space_after = Pt(0)
            except Exception:
                pass

            # 組織另起一行
            if org_val:
                org_p = c1.add_paragraph()
                # 清空 runs（通常新段落沒 runs）
                org_p.text = org_val
                if org_p.runs:
                    set_run_font(org_p.runs[0], font_pt, bold=False)
                try:
                    org_p.paragraph_format.space_before = Pt(0)
                    org_p.paragraph_format.space_after = Pt(0)
                except Exception:
                    pass

            # 垂直置中第二欄
            try:
                c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass

            # ---------- 第三欄：簽到欄（NBSP 占位） ----------
            c2 = row_cells[2]
            p2 = c2.paragraphs[0] if c2.paragraphs else c2.add_paragraph()
            # 清除 runs 並放 NBSP
            for r in list(p2.runs):
                try:
                    r._element.getparent().remove(r._element)
                except Exception:
                    pass
            r2 = p2.add_run("\u00A0")
            try:
                r2.font.size = Pt(font_pt)
            except Exception:
                pass
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            try:
                c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                pass

    p_e = doc.add_paragraph()
    p_e.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_e_run = p_e.add_run(slash_date)
    set_run_font(p_e_run, FONT_PT, bold=False)

    # 儲存
    doc.save(out_path)
    print(f"Saved sign-in sheet to {out_path} (page_width_cm={page_width_cm:.2f}, available_cm={available_cm:.2f}, cols_cm={[col0, col1, col2]})")

if __name__ == "__main__":
    main()
