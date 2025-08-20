#!/usr/bin/env python3
"""
Send personalized email(s) for an identifier or for all followers of a program.

Usage examples:
  # send for a single identifier (keeps old behavior)
  python scripts/actions/send_email_with_attachments.py jam123 --send

  # send to all followers of program id 5 using an explicit follower file
  python scripts/actions/send_email_with_attachments.py --program 5 --followers-file data/follower.xlsx --sheet-name "Sheet1" --draft

Notes:
 - Template rendering uses Jinja2 if installed (recommended). Otherwise falls back to a simple replace.
 - Supports JSON, .xlsx, .xls (xls requires xlrd).
"""
from __future__ import annotations

import argparse
import importlib
import json
import logging
import mimetypes
import os
import re
from email.message import EmailMessage
from pathlib import Path
from typing import Iterable, List, Dict, Any, Optional
import smtplib
import html as _html_mod

# optional deps
try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

try:
    import jinja2
except ModuleNotFoundError:
    jinja2 = None

# main logging (set INFO by default; set to DEBUG if you need more)
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# --- Try to import project template_utils; fall back to None if unavailable ---
try:
    from scripts.actions import template_utils as tu
    tu_render_body = getattr(tu, "render_body_from_template", None)
    tu_render_docx = getattr(tu, "render_docx_template", None)
except Exception:
    tu = None
    tu_render_body = None
    tu_render_docx = None

import tempfile
import uuid
# --- end import ---

# -------------------- DEFAULT PATHS (adjust if needed) --------------------
BASE_DIR = Path(r"C:\Users\User\activity_management")
DEFAULT_DATA_DIR = BASE_DIR / "data"
DEFAULT_SHARED_JSON = DEFAULT_DATA_DIR / "shared" / "program_data.json"
DEFAULT_TEMPLATE = BASE_DIR / "templates" / "letter_sample" / "活動參與通知.docx"
DEFAULT_ATTACHMENTS_DIR = DEFAULT_DATA_DIR / "attachments"
# -------------------------------------------------------------------------

# -------------------- SMTP config loader ---------------------------------
def load_smtp_config(path: Path) -> None:
    try:
        if not path.exists():
            return
        with path.open("r", encoding="utf-8") as fh:
            conf = json.load(fh)
        if isinstance(conf, dict):
            if conf.get("smtp_server") and not os.environ.get("SMTP_SERVER"):
                os.environ["SMTP_SERVER"] = str(conf["smtp_server"])
            if conf.get("smtp_port") and not os.environ.get("SMTP_PORT"):
                os.environ["SMTP_PORT"] = str(conf["smtp_port"])
            if conf.get("smtp_username") and not os.environ.get("SMTP_USERNAME"):
                os.environ["SMTP_USERNAME"] = str(conf["smtp_username"])
            if conf.get("smtp_password") and not os.environ.get("SMTP_PASSWORD"):
                os.environ["SMTP_PASSWORD"] = str(conf["smtp_password"])
    except Exception as e:
        logging.warning("Failed to load smtp config %s: %s", path, e)


# -------------------- program_data loader & matcher -----------------------
def load_programs(path: Path) -> List[Dict[str, Any]]:
    if not path.exists():
        return []
    try:
        with path.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
    except Exception:
        logging.warning("Failed to load program data from %s", path)
        return []
    if isinstance(data, dict):
        return [data]
    if isinstance(data, list):
        return data
    return []


def find_program_by_id(programs: List[Dict[str, Any]], pid: str) -> Optional[Dict[str, Any]]:
    if not programs:
        return None
    pid_s = str(pid).strip().lower()
    for prog in programs:
        if str(prog.get("id", "")).strip().lower() == pid_s:
            return prog
        # also check planName/title
        if str(prog.get("planName", "")).strip().lower() == pid_s:
            return prog
        if str(prog.get("plan_name", "")).strip().lower() == pid_s:
            return prog
    return None


# -------------------- data file discovery & loading -----------------------
def find_data_file_by_id(data_dir: Path, target_id: str) -> Optional[Path]:
    """Search data_dir (recursive) for JSON or Excel that contains target_id."""
    target = str(target_id).strip()
    # JSON first
    for p in sorted(data_dir.rglob("*.json")):
        try:
            # skip shared program_data file
            if p.resolve() == DEFAULT_SHARED_JSON.resolve():
                continue
            with p.open("r", encoding="utf-8") as fh:
                data = json.load(fh)
        except Exception:
            continue
        if isinstance(data, dict):
            if any(str(k).strip() == target for k in data.keys()):
                return p
            if str(data.get("id", "")).strip() == target:
                return p
            for v in data.values():
                if isinstance(v, list):
                    for item in v:
                        try:
                            if str(item.get("id", "")).strip() == target:
                                return p
                        except Exception:
                            continue
        elif isinstance(data, list):
            for item in data:
                try:
                    if str(item.get("id", "")).strip() == target:
                        return p
                except Exception:
                    continue
    # Excel
    try:
        openpyxl = importlib.import_module("openpyxl")
    except ModuleNotFoundError:
        openpyxl = None
    try:
        xlrd = importlib.import_module("xlrd")
    except ModuleNotFoundError:
        xlrd = None

    for p in sorted(data_dir.rglob("*.xls")) + sorted(data_dir.rglob("*.xlsx")):
        if p.resolve() == DEFAULT_SHARED_JSON.resolve():
            continue
        ext = p.suffix.lower()
        if ext == ".xlsx" and openpyxl:
            try:
                wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
            except Exception:
                continue
            for sname in wb.sheetnames:
                ws = wb[sname]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                headers = [str(h).strip().lower() if h is not None else "" for h in rows[0]]
                id_idx = [i for i, h in enumerate(headers) if "id" in h and h != ""]
                if not id_idx:
                    continue
                for row in rows[1:]:
                    for idx in id_idx:
                        if idx >= len(row):
                            continue
                        val = row[idx]
                        if val is None:
                            continue
                        if str(val).strip() == target:
                            wb.close()
                            return p
            wb.close()
        elif ext == ".xls" and xlrd:
            try:
                book = xlrd.open_workbook(str(p))
            except Exception:
                continue
            for sname in book.sheet_names():
                sh = book.sheet_by_name(sname)
                if sh.nrows == 0:
                    continue
                headers = [str(sh.cell_value(0, c)).strip().lower() for c in range(sh.ncols)]
                id_idx = [i for i, h in enumerate(headers) if "id" in h and h != ""]
                if not id_idx:
                    continue
                for r in range(1, sh.nrows):
                    for idx in id_idx:
                        try:
                            val = sh.cell_value(r, idx)
                        except Exception:
                            continue
                        if val is None or str(val).strip() == "":
                            continue
                        if str(val).strip() == target:
                            return p
    return None


def load_records(path: Path, sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
    """Load records from JSON, .xlsx, or .xls. Return list of dicts (keys lowercased)."""
    ext = path.suffix.lower()
    if ext == ".json":
        with path.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, dict):
            # dict of objects
            if all(isinstance(v, dict) for v in data.values()):
                out = []
                for k, v in data.items():
                    rec = {str(kk).strip().lower(): vv for kk, vv in v.items()}
                    if "id" not in rec:
                        rec["id"] = str(k)
                    out.append(rec)
                return out
            data = [data]
        return [{str(k).strip().lower(): v for k, v in r.items()} for r in data]

    if ext in {".xlsx", ".xls"}:
        # prefer openpyxl for xlsx, xlrd for xls
        if ext == ".xlsx":
            openpyxl = importlib.import_module("openpyxl")
            wb = openpyxl.load_workbook(path, data_only=True)
            if sheet_name:
                if sheet_name not in wb.sheetnames:
                    raise ValueError(f"Sheet '{sheet_name}' not found in {path} (available: {wb.sheetnames})")
                ws = wb[sheet_name]
            else:
                ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return []
            headers = [str(h).strip().lower() if h is not None else "" for h in rows[0]]
            recs: List[Dict[str, Any]] = []
            for row in rows[1:]:
                r = {}
                for i, val in enumerate(row):
                    header = headers[i] if i < len(headers) else f"col_{i}"
                    r[header] = val
                recs.append(r)
            return recs
        else:
            # .xls using xlrd
            xlrd = importlib.import_module("xlrd")
            book = xlrd.open_workbook(str(path))
            if sheet_name:
                if sheet_name not in book.sheet_names():
                    raise ValueError(f"Sheet '{sheet_name}' not found in {path} (available: {book.sheet_names()})")
                sh = book.sheet_by_name(sheet_name)
            else:
                sh = book.sheet_by_index(0)
            if sh.nrows == 0:
                return []
            headers = [str(sh.cell_value(0, c)).strip().lower() for c in range(sh.ncols)]
            recs = []
            for r in range(1, sh.nrows):
                rowvals = [sh.cell_value(r, c) for c in range(sh.ncols)]
                rec = {}
                for i, val in enumerate(rowvals):
                    header = headers[i] if i < len(headers) else f"col_{i}"
                    rec[header] = val
                recs.append(rec)
            return recs

    raise ValueError(f"Unsupported file extension: {ext}")


def load_all_records_from_dir(data_dir: Path, sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
    """Load all records from JSON/xlsx/xls files under data_dir except files in shared/."""
    out: List[Dict[str, Any]] = []
    for p in sorted(data_dir.rglob("*")):
        if not p.is_file():
            continue
        if p.resolve() == DEFAULT_SHARED_JSON.resolve():
            continue
        if "shared" in p.parts:
            # skip shared folder files (except if you want to include)
            continue
        if p.suffix.lower() in {".json", ".xlsx", ".xls"}:
            try:
                recs = load_records(p, sheet_name=sheet_name)
                out.extend(recs)
            except Exception:
                logging.debug("Failed to load %s, skipping.", p)
                continue
    return out


# -------------------- template rendering -----------------------------------
# Use project template_utils.render_body_from_template if available; otherwise provide
# a small inline fallback that does a simple paragraph-level replacement.
def _fallback_render_body_from_template(template_path: Path, context: Dict[str, Any]) -> str:
    """
    Simple fallback: paragraph-level replacement (no jinja2, no deep flatten).
    Use this only if project template_utils is not available.
    """
    if Document is None:
        raise ModuleNotFoundError("python-docx is required for template rendering (fallback).")
    doc = Document(str(template_path))

    def _flatten(prefix: str, val: Any, out: Dict[str, str]):
        if isinstance(val, dict):
            for k, v in val.items():
                _flatten(f"{prefix}.{k}" if prefix else str(k), v, out)
        elif isinstance(val, (list, tuple)):
            for i, item in enumerate(val):
                _flatten(f"{prefix}[{i}]" if prefix else f"[{i}]", item, out)
        else:
            out[prefix] = "" if val is None else str(val)

    flat: Dict[str, str] = {}
    for k, v in context.items():
        _flatten(str(k), v, flat)

    placeholder_re = re.compile(r"\{\{\s*(.*?)\s*\}\}")

    # Embedded minimal helpers. These only cover common cases needed for
    # fallback rendering when project-level template_utils is unavailable.

    def format_chinese_date(value: Any) -> str:
        """Basic YYYY年M月D日(星期X) formatter for ISO-like strings."""
        from datetime import datetime, date
        if value is None:
            return ""
        if isinstance(value, datetime):
            dt = value.date()
        elif isinstance(value, date):
            dt = value
        else:
            try:
                dt = datetime.fromisoformat(str(value)).date()
            except Exception:
                return str(value)
        wmap = ["一", "二", "三", "四", "五", "六", "日"]
        return f"{dt.year}年{dt.month}月{dt.day}日(星期{wmap[dt.weekday()]})"

    def _wrap_highlight(s: str) -> str:
        """No-op highlight wrapper for fallback rendering."""
        return s

    def _apply_filters(val: Any, filters: List[str]) -> str:
        s = "" if val is None else str(val)
        for f in filters:
            if tu is None:
                if f in ("cn_date", "cnDate", "format_date"):
                    s = format_chinese_date(s)
                elif f in ("hl", "highlight"):
                    s = _wrap_highlight(s)
            else:
                if f in ("cn_date", "cnDate", "format_date") and hasattr(tu, "format_chinese_date"):
                    s = tu.format_chinese_date(s)
                elif f in ("hl", "highlight") and hasattr(tu, "_wrap_highlight"):
                    s = tu._wrap_highlight(s)
        return s

    def render_text(text: str) -> str:
        if not text:
            return text

        def _replace(m: re.Match) -> str:
            expr = m.group(1).strip()
            parts = [p.strip() for p in expr.split("|") if p.strip()]
            if not parts:
                return m.group(0)
            key = re.sub(r"\s+", "", parts[0])
            filters = parts[1:]
            val = flat.get(key)
            if val is None:
                return m.group(0)
            return _apply_filters(val, filters)

        return placeholder_re.sub(_replace, text)
    def process_paragraphs(paragraphs: Iterable[Any]):
        for para in paragraphs:
            if "{{" in para.text or "{%" in para.text:
                try:
                    new_text = render_text(para.text)
                except Exception:
                    new_text = para.text
                if new_text is not None:
                    new_text = new_text.replace("\\n", "\n")
                if para.runs:
                    para.runs[0].text = new_text
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(new_text)

    # process top-level paragraphs
    process_paragraphs(doc.paragraphs)

    # process paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    body_lines: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                body_lines.extend(p.text for p in cell.paragraphs)
    return "\n".join(body_lines).strip()


def render_body_from_template(template_path: Path, context: Dict[str, Any]) -> str:
    """
    Prefer project-provided template util if available; otherwise use fallback.
    """
    if tu_render_body is not None:
        try:
            return tu_render_body(template_path, context)
        except Exception as e:
            logging.debug("tu.render_body_from_template failed: %s; falling back.", e)
    return _fallback_render_body_from_template(template_path, context)


# -------------------- attachments & helpers --------------------------------
def sanitize_filename(s: str, max_len: int = 200) -> str:
    s = re.sub(r"[\\/:\*\?\"<>\|]+", "-", s or "")
    s = re.sub(r"\s+", " ", s).strip()
    return s[:max_len]


def attach_file_to_msg(msg: EmailMessage, p: Path):
    if not p.exists():
        logging.debug("Attachment not found: %s", p)
        return
    ctype, _ = mimetypes.guess_type(str(p))
    if ctype:
        maintype, subtype = ctype.split("/", 1)
    else:
        maintype, subtype = "application", "octet-stream"
    with p.open("rb") as fh:
        data = fh.read()
    msg.add_attachment(data, maintype=maintype, subtype=subtype, filename=p.name)


def attach_entries_from_list(msg: EmailMessage, entries: List[str], include_pdfs: bool = True):
    for e in entries:
        if not e:
            continue
        p = Path(e)
        if not p.is_absolute():
            p = (BASE_DIR / e).resolve()
        if p.is_dir():
            for f in sorted(p.glob("*")):
                if not f.is_file():
                    continue
                if not include_pdfs and f.suffix.lower() == ".pdf":
                    continue
                attach_file_to_msg(msg, f)
        else:
            attach_file_to_msg(msg, p)


# -------------------- message creation -------------------------------------
def create_message(record: Dict[str, Any], template_path: Optional[Path], attachments_entries: List[str], include_pdfs: bool, templates_dir: Optional[Path]) -> EmailMessage:
    msg = EmailMessage()

    # extended list of possible email field names (English + common Chinese variants)
    EMAIL_FIELDS = [
        "mail", "email", "to", "recipient", "e-mail", "email_address", "address", "contact_email",
        "電子郵件", "電子郵箱", "信箱", "聯絡信箱", "emailaddress"
    ]
    to_field = ""
    for ef in EMAIL_FIELDS:
        # keys in record are lowercased by load_records; chinese also present as-is but lowercased
        val = record.get(ef)
        if val:
            to_field = str(val).strip()
            break
    if not to_field:
        # last resort: check any value that looks like an email using regex
        email_re = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
        for k, v in record.items():
            try:
                if v and isinstance(v, str) and email_re.search(v):
                    to_field = v.strip()
                    break
            except Exception:
                continue
    if not to_field:
        logging.warning("No recipient email found in record. Checked fields: %s. Record keys: %s", EMAIL_FIELDS, list(record.keys()))

    # Subject composition: prefer record.subject else template_name - program.planName - event0 - date
    if record.get("subject"):
        subject = str(record.get("subject"))
    else:
        tpl_name = Path(template_path).stem if template_path else ""
        prog = record.get("program_data") or {}
        plan_name = prog.get("planName") or prog.get("plan_name") or prog.get("plan") or ""

        event0 = ""
        possible_event_keys = ("eventNames", "event_names", "events", "eventList", "event_list")
        for k in possible_event_keys:
            if k in prog:
                ev = prog.get(k)
                if isinstance(ev, (list, tuple)) and ev:
                    event0 = str(ev[0])
                elif isinstance(ev, str) and ev:
                    event0 = ev
                break

        date_keys = ("date", "start_date", "event_date", "dates")
        date_str = ""
        for dk in date_keys:
            if dk in prog and prog.get(dk):
                dval = prog.get(dk)
                if isinstance(dval, (list, tuple)) and dval:
                    dval = dval[0]
                date_str = str(dval).strip()
                break
        if not date_str:
            from datetime import datetime
            date_str = datetime.now().strftime("%Y-%m-%d")

        parts = []
        if tpl_name:
            parts.append(f"【{tpl_name}】")
        if date_str:
            parts.append(date_str)
        if plan_name:
            parts.append(plan_name)
        if event0:
            parts.append(event0)
        subject = " - ".join(parts) if parts else "活動通知"

    from_addr = os.environ.get("SMTP_USERNAME", "noreply@example.com")

    if to_field:
        msg["To"] = to_field
    if record.get("cc"):
        msg["Cc"] = str(record.get("cc"))
    if record.get("bcc"):
        msg["Bcc"] = str(record.get("bcc"))
    msg["Subject"] = subject
    msg["From"] = from_addr

    body = record.get("body", "") or ""
    if template_path and template_path.exists():
        context = {}
        if "follower" in record and isinstance(record["follower"], dict):
            context["follower"] = {k: ("" if v is None else v) for k, v in record["follower"].items()}
        else:
            context["follower"] = {k: ("" if v is None else v) for k, v in record.items()}
        if "program_data" in record and isinstance(record["program_data"], dict):
            context["program_data"] = record["program_data"]
        else:
            context["program_data"] = record.get("program_data", {}) or {}
        try:
            body = render_body_from_template(template_path, context)
        except Exception as e:
            logging.warning("Template render failed for %s: %s", template_path, e)

    # ===== 保險：把 literal "\n" 轉為真正換行，並在 debug 下顯示替換前後 repr =====
    if body:
        logging.debug("Rendered body (repr, first 1000 chars): %r", body[:1000])
        body = body.replace("\\n", "\n")
        logging.debug("After replace (repr, first 1000 chars): %r", body[:1000])
    # =======================================================================

    # html handling: if record provides html, use it; otherwise auto-generate a simple html from body
    html = record.get("html")
    if html:
        msg.set_content(body if body else "This message contains HTML content.")
        msg.add_alternative(str(html), subtype="html")
    else:
        msg.set_content(body)
        try:
            safe = _html_mod.escape(body or "")
            html_generated = safe.replace("\n", "<br/>")
            msg.add_alternative(html_generated, subtype="html")
        except Exception:
            logging.debug("Failed to add html alternative; continuing with plain text only.")

    attach_entries_from_list(msg, attachments_entries or [], include_pdfs=include_pdfs)
    if templates_dir:
        td = Path(templates_dir)
        if td.exists():
            for f in sorted(td.glob("*")):
                if f.suffix.lower() in {".docx", ".doc"}:
                    attach_file_to_msg(msg, f)
    return msg


# -------------------- send & save ------------------------------------------
def send_all_messages(messages: List[EmailMessage]) -> None:
    server = os.environ.get("SMTP_SERVER")
    if not server:
        raise KeyError("SMTP_SERVER environment variable is required to send emails.")
    port = int(os.environ.get("SMTP_PORT", 587))
    username = os.environ.get("SMTP_USERNAME")
    password = os.environ.get("SMTP_PASSWORD")
    if not (username and password):
        raise KeyError("SMTP_USERNAME and SMTP_PASSWORD must be set to send emails.")

    with smtplib.SMTP(server, port) as smtp:
        smtp.starttls()
        smtp.login(username, password)
        for msg in messages:
            recipients = []
            for hdr in ("To", "Cc", "Bcc"):
                val = msg.get(hdr)
                if val:
                    parts = re.split(r"[,;]+", val)
                    recipients.extend([p.strip() for p in parts if p.strip()])
            logging.info("Sending to %s (subject: %s)", recipients, msg.get("Subject"))
            smtp.send_message(msg, from_addr=msg.get("From"), to_addrs=recipients)


def save_draft(msg: EmailMessage, directory: Path) -> None:
    directory.mkdir(parents=True, exist_ok=True)
    to_safe = sanitize_filename(msg.get("To", "unknown"))
    subj_safe = sanitize_filename(msg.get("Subject", "no-subject"))
    filename = f"{to_safe}_{subj_safe}.eml"
    path = directory / filename
    i = 1
    while path.exists():
        path = directory / f"{to_safe}_{subj_safe}-{i}.eml"
        i += 1
    with path.open("wb") as fh:
        fh.write(msg.as_bytes())
    logging.info("Saved draft: %s", path)


# -------------------- helper: determine if record belongs to program ---------
def record_matches_program(record: Dict[str, Any], program: Dict[str, Any]) -> bool:
    """Return True if record references the given program (by id or name)."""
    if not program:
        return False
    pid = str(program.get("id", "")).strip().lower()
    pname = str(program.get("planName", "") or program.get("plan_name", "")).strip().lower()
    # check common fields in record
    candidate_keys = ("planid", "plan_id", "activity_id", "activityid", "program_id", "programid", "id", "planname", "plan_name", "program", "plan")
    for k in candidate_keys:
        v = record.get(k)
        if v is None:
            continue
        vs = str(v).strip().lower()
        if vs == pid or vs == pname or (pname and pname in vs):
            return True
        # if record field is a list or comma separated values, check items
        if isinstance(v, (list, tuple)):
            for it in v:
                if str(it).strip().lower() in (pid, pname):
                    return True
        if isinstance(v, str) and "," in v:
            for it in [x.strip().lower() for x in v.split(",") if x.strip()]:
                if it in (pid, pname):
                    return True
    return False


# -------------------- main --------------------------------------------------
def main(argv: Optional[Iterable[str]] = None) -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("identifier", nargs="?", default=None, help="ID to locate the data record in data/ (optional when --program used)")
    parser.add_argument("--program", type=str, default=None, help="Program id to send to all followers of that program (e.g. 5)")
    parser.add_argument("--followers-file", type=Path, default=None, help="Optional follower file (json/xlsx/xls) to use when --program provided; if omitted, script will search data dir and match by program id as before.")
    parser.add_argument("--data-dir", type=Path, default=DEFAULT_DATA_DIR, help="Directory to search for data files")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help="Template docx to render body (optional)")
    parser.add_argument("--attachments-dir", type=Path, default=DEFAULT_ATTACHMENTS_DIR, help="Default attachments dir")
    parser.add_argument("--sheet-name", type=str, default=None, help="Excel sheet name to read (optional)")
    parser.add_argument("--templates-folder", type=Path, default=None, help="Folder whose .docx/.doc to attach in addition")
    parser.add_argument("--output", type=Path, default=Path("output/drafts"), help="Directory to save drafts")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--send", action="store_true", help="Send emails via SMTP")
    group.add_argument("--draft", action="store_true", help="Save drafts only (default)")
    args = parser.parse_args(argv)

    if not args.send and not args.draft:
        args.draft = True

    load_smtp_config(Path("config/smtp.json"))

    # load programs
    programs = load_programs(DEFAULT_SHARED_JSON)

    messages: List[EmailMessage] = []

    # -------- PROGRAM MODE: gather all follower records that belong to the program -------
    if args.program:
        prog = find_program_by_id(programs, args.program)
        if not prog:
            logging.error("Program with id '%s' not found in %s", args.program, DEFAULT_SHARED_JSON)
            return
        logging.info("Operating in program mode for program id %s -> %s", args.program, prog.get("planName") or prog.get("id"))

        # If user provided an explicit follower file, load it and use all its records
        follower_records: List[Dict[str, Any]] = []
        if args.followers_file:
            ff = Path(args.followers_file)
            if not ff.exists():
                logging.error("Followers file %s not found.", ff)
                return
            try:
                follower_records = load_records(ff, sheet_name=args.sheet_name)
            except Exception as e:
                logging.error("Failed to load followers from %s: %s", ff, e)
                return
            if not follower_records:
                logging.error("No follower records loaded from %s", ff)
                return
            logging.info("Loaded %d follower records from %s", len(follower_records), ff)
        else:
            # original behavior: scan data dir and filter those that match program
            all_recs = load_all_records_from_dir(args.data_dir, sheet_name=args.sheet_name)
            if not all_recs:
                logging.error("No records found under %s", args.data_dir)
                return
            follower_records = [r for r in all_recs if record_matches_program(r, prog)]
            if not follower_records:
                logging.error(
                    "No follower records matched program id %s under %s. "
                    "If your follower file does not contain program identifiers, "
                    "consider using --followers-file to point directly to the follower list.",
                    args.program, args.data_dir
                )
                return

        # create messages for each follower record and attach program data
        for rec in follower_records:
            rec_for_message = dict(rec)
            rec_for_message["program_data"] = prog
            rec_for_message["follower"] = {k: ("" if v is None else v) for k, v in rec.items()}

            # attachments resolution (same logic as before)
            attachments_entries: List[str] = []
            for candidate_field in ("attach_column", "attachments", "attachment", "attachments_dir", "attachment_path"):
                val = rec.get(candidate_field)
                if val:
                    parts = [p.strip() for p in str(val).split(",") if p.strip()]
                    attachments_entries.extend(parts)
                    break
            if not attachments_entries:
                candidate_folder = Path(args.attachments_dir) / str(rec.get("id") or rec.get("to") or "").strip()
                if candidate_folder.exists():
                    attachments_entries.append(str(candidate_folder))
                else:
                    common = Path(args.attachments_dir) / "common"
                    if common.exists():
                        attachments_entries.append(str(common))

            # template override support
            body_template_path = None
            if rec.get("body_template"):
                body_template_path = Path(str(rec.get("body_template")))
            elif prog and isinstance(prog, dict) and prog.get("body_template"):
                body_template_path = Path(str(prog.get("body_template")))
            else:
                body_template_path = args.template

            try:
                msg = create_message(
                    rec_for_message,
                    template_path=Path(body_template_path) if body_template_path else None,
                    attachments_entries=attachments_entries,
                    include_pdfs=True,
                    templates_dir=args.templates_folder,
                )
                messages.append(msg)
            except Exception as e:
                logging.error("Failed to create message for record %s: %s", rec, e)

    # -------- IDENTIFIER MODE: behave as previous (single record lookup) -----------
    else:
        if not args.identifier:
            logging.error("No identifier provided. Either pass an identifier or use --program <id>.")
            return
        found = find_data_file_by_id(args.data_dir, args.identifier)
        if not found:
            logging.error("No data file containing ID %s found under %s", args.identifier, args.data_dir)
            return
        logging.info("Using data file: %s", found)
        records = load_records(found, sheet_name=args.sheet_name)
        if not records:
            logging.error("No records loaded from %s", found)
            return

        def matches(rec: Dict[str, Any], ident: str) -> bool:
            for k in ("id", "to", "recipient", "email"):
                v = rec.get(k)
                if v is None:
                    continue
                if str(v).strip().lower() == ident.strip().lower():
                    return True
            return False

        matched = [r for r in records if matches(r, args.identifier)]
        if not matched:
            logging.error("No record matching identifier %s inside %s", args.identifier, found)
            return

        for rec in matched:
            # fixed: initialize prog properly (no walrus misuse)
            prog = None
            # try to find program for this record from loaded programs
            for p in programs:
                if record_matches_program(rec, p):
                    prog = p
                    break

            rec_for_message = dict(rec)
            rec_for_message["program_data"] = prog or {}
            rec_for_message["follower"] = {k: ("" if v is None else v) for k, v in rec.items()}

            attachments_entries: List[str] = []
            for candidate_field in ("attach_column", "attachments", "attachment", "attachments_dir", "attachment_path"):
                val = rec.get(candidate_field)
                if val:
                    parts = [p.strip() for p in str(val).split(",") if p.strip()]
                    attachments_entries.extend(parts)
                    break
            if not attachments_entries:
                candidate_folder = Path(args.attachments_dir) / str(args.identifier).strip()
                if candidate_folder.exists():
                    attachments_entries.append(str(candidate_folder))
                else:
                    common = Path(args.attachments_dir) / "common"
                    if common.exists():
                        attachments_entries.append(str(common))

            body_template_path = None
            if rec.get("body_template"):
                body_template_path = Path(str(rec.get("body_template")))
            elif prog and isinstance(prog, dict) and prog.get("body_template"):
                body_template_path = Path(str(prog.get("body_template")))
            else:
                body_template_path = args.template

            try:
                msg = create_message(
                    rec_for_message,
                    template_path=Path(body_template_path) if body_template_path else None,
                    attachments_entries=attachments_entries,
                    include_pdfs=True,
                    templates_dir=args.templates_folder,
                )
                messages.append(msg)
            except Exception as e:
                logging.error("Failed to create message for record %s: %s", rec, e)

    # ---- send or save drafts ----
    if args.send:
        try:
            send_all_messages(messages)
        except Exception as e:
            logging.error("Failed to send messages: %s", e)
    else:
        for m in messages:
            try:
                save_draft(m, args.output)
            except Exception as e:
                logging.error("Failed to save draft for %s: %s", m.get("To"), e)


if __name__ == "__main__":
    main()
