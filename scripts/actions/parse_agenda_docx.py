# parse_agenda_docx.py
# pip install python-docx
from __future__ import annotations

import json
import re
from pathlib import Path
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple
from collections import Counter

from docx import Document

TIME_RE = re.compile(r"(\d{1,2}:\d{2})\s*-\s*(\d{1,2}:\d{2})")

def _clean_text(s: str) -> str:
    return (s or "").replace("\u00A0", " ").strip()

def _extract_time_range(text: str) -> Optional[Tuple[str, str]]:
    if not text:
        return None
    for line in text.splitlines():
        m = TIME_RE.search(line)
        if m:
            return (m.group(1), m.group(2))
    m = TIME_RE.search(text)
    if m:
        return (m.group(1), m.group(2))
    return None

def _parse_time_to_dt(hhmm: str) -> datetime:
    return datetime.strptime(hhmm, "%H:%M")

def _minutes_between(a: str, b: str) -> int:
    dt_a, dt_b = _parse_time_to_dt(a), _parse_time_to_dt(b)
    if dt_b < dt_a:
        dt_b += timedelta(days=1)
    return int((dt_b - dt_a).total_seconds() // 60)

def _pick_agenda_table(doc: Document):
    """Pick the first table that looks like a 3-column agenda (時間/主題/講者).
    選擇第一個看起來像「時間/主題/講者」的表格。"""
    def norm(s: str) -> str:
        return _clean_text(s).replace(" ", "").replace("\n", "")
    for tbl in doc.tables:
        if len(tbl.columns) >= 3 and len(tbl.rows) >= 2:
            headers = []
            for i in range(min(3, len(tbl.rows))):
                cells = tbl.rows[i].cells
                if len(cells) >= 3:
                    headers.append((norm(cells[0].text), norm(cells[1].text), norm(cells[2].text)))
            header_text = " ".join(["|".join(h) for h in headers])
            if ("時間" in header_text) and ("主題" in header_text or "主旨" in header_text or "議題" in header_text) and ("講者" in header_text or "講師" in header_text):
                return tbl
    return doc.tables[0] if doc.tables else None

# ---- Special 正規化規則（同義字 → 統一標題）----
SPECIAL_RULES: List[Tuple[re.Pattern, str, bool]] = [
    # (pattern, normalized_title, is_end_of_day)
    (re.compile(r"(問卷|後測|後測與問卷|問卷與後測|closing|結語)", re.I), "問卷與後測", True),
    (re.compile(r"(綜合討論|討論|討論時間)", re.I), "綜合討論", False),
    (re.compile(r"(兌獎)", re.I), "兌獎", False),
    (re.compile(r"(午餐|午休|午間)", re.I), "午間休息", False),
    (re.compile(r"(大合照|合照)", re.I), "大合照", False),
    (re.compile(r"(致詞|開場)", re.I), "致詞", False),
    # 將 "主持" 與 "主持人" 同一化為 "主持人"
    (re.compile(r"(主持|主持人)", re.I), "主持人", False),
    (re.compile(r"(休息|中場休息)", re.I), "課間休息", False),
    (re.compile(r"(報到|簽到|簽退)", re.I), "報到／簽到簽退", False),
]
END_TITLES = { title for _, title, is_end in SPECIAL_RULES if is_end }

# 偵測主持人的 pattern（中英皆列）
HOST_PATTERN = re.compile(r"(?:主持人|主持：|主持|主持者|Chair|Moderator|主持人：)\s*[:：]?\s*(?P<name>[^,\n;（）()]*)", re.I)

def _match_special_title(topic_text: str, speaker_text: str) -> Tuple[Optional[str], bool]:
    """回傳 (normalized_title, is_end_of_day)；若無匹配則 (None, False)"""
    hay = "{}\n{}".format(_clean_text(topic_text), _clean_text(speaker_text))
    for pat, norm_title, is_end in SPECIAL_RULES:
        if pat.search(hay):
            return norm_title, is_end
    return None, False

def _extract_host_names(speaker_text: str) -> List[str]:
    """從講者欄位抓出主持人名稱（可能有多個），回傳名稱列表（已清理）"""
    if not speaker_text:
        return []
    names = []
    for m in HOST_PATTERN.finditer(speaker_text):
        name = m.groupdict().get("name") or ""
        name = name.strip()
        if not name:
            cleaned = re.sub(r"(主持人|主持：|主持|主持者|Chair|Moderator)", "", speaker_text, flags=re.I).strip()
            if cleaned:
                name = cleaned.split("\n")[0].strip()
        if name:
            names.append(name)
    return names

# ----- New helpers to avoid duplicate host/special entries -----
def _normalize_speaker_for_compare(s: Optional[str]) -> str:
    """去掉主持字樣並小寫化，用於比對相等性/包含關係"""
    if not s:
        return ""
    cleaned = re.sub(r"(主持人|主持：|主持|主持者|Chair|Moderator|主持人：)", "", s, flags=re.I)
    cleaned = re.sub(r"[:：,，;；\(\)（）\"']", " ", cleaned)
    return " ".join(cleaned.split()).strip().lower()

def _special_exists(specials: List[Dict[str, Any]], title: str, speaker_norm: str) -> bool:
    """
    檢查 specials 中是否已有相同 title 並且 speaker（正規化後）相同或包含 speaker_norm。
    如果 speaker_norm 為空，檢查是否已有相同 title 的 special（無 speaker）。
    """
    for sp in specials:
        if sp.get("title") != title:
            continue
        sp_s = sp.get("speaker", "")
        sp_norm = _normalize_speaker_for_compare(sp_s)
        if speaker_norm:
            if speaker_norm in sp_norm or sp_norm in speaker_norm:
                return True
        else:
            if not sp_s:
                return True
    return False

def _add_special(specials: List[Dict[str, Any]], special_times: List[Optional[Tuple[str, str]]],
                 title: str, after_speaker: int, duration: Optional[int], speaker: Optional[str],
                 time_rng: Optional[Tuple[str,str]]) -> Optional[int]:
    """
    統一加 special，加入前做重複檢查避免同一人重複出現於主持/主持人/主持等 title。
    如果成功加入，回傳該 special 在 specials 中的 index；若已存在則回傳 None。
    """
    speaker_norm = _normalize_speaker_for_compare(speaker)
    if _special_exists(specials, title, speaker_norm):
        return None  # 已存在，不重複加入
    obj: Dict[str, Any] = {
        "after_speaker": after_speaker,
        "title": title,
        "duration": duration
    }
    if speaker:
        obj["speaker"] = speaker
    specials.append(obj)
    special_times.append(time_rng)
    return len(specials) - 1

# -----------------------------------------------------------------

def parse_agenda(docx_path: Path, event_name: str) -> Dict[str, Any]:
    """Parse the docx and return a program dict with agenda_settings + speakers.
    解析 docx，回傳含 agenda_settings 與 speakers 的 program dict。"""
    doc = Document(docx_path)
    tbl = _pick_agenda_table(doc)
    if tbl is None:
        raise SystemExit("找不到任何表格。請確認檔案內含有 3 欄的議程表。")

    # 如果第一列是表頭（時間/主題/講者），就跳過
    start_row_idx = 0
    if len(tbl.rows) >= 1:
        heads = [_clean_text(c.text) for c in tbl.rows[0].cells[:3]]
        if any("時間" in h for h in heads) and any(("主題" in h or "主旨" in h) for h in heads) and any(("講者" in h or "講師" in h) for h in heads):
            start_row_idx = 1

    talks: List[Dict[str, Any]] = []
    specials: List[Dict[str, Any]] = []
    special_times: List[Optional[Tuple[str, str]]] = []  # 與 specials 對齊
    timeline: List[Tuple[str, int]] = []  # ("talk", talk_no) or ("special", idx)

    first_time_start: Optional[str] = None
    last_time_end: Optional[str] = None
    talk_idx = 0
    talk_durations: List[int] = []

    for r in tbl.rows[start_row_idx:]:
        cells = r.cells
        if len(cells) < 3:
            # 不規則列，忽略
            continue

        time_text = _clean_text(cells[0].text)
        topic_text = _clean_text(cells[1].text)
        speaker_text = _clean_text(cells[2].text)

        time_rng = _extract_time_range(time_text)
        norm_title, is_end_flag = _match_special_title(topic_text, speaker_text)

        # 若是 special（例如「致詞」或「主持」），建立 special 並嘗試帶上 speaker 資訊
        if norm_title is not None:
            duration_val: Optional[int] = None
            if time_rng is not None:
                duration_val = _minutes_between(time_rng[0], time_rng[1])
                if first_time_start is None:
                    first_time_start = time_rng[0]
                last_time_end = time_rng[1]

            # 先把 main special 加入（若尚未存在）
            sidx = _add_special(specials, special_times, norm_title, talk_idx if talk_idx > 0 else 0, duration_val, speaker_text or None, time_rng)
            if sidx is not None:
                timeline.append(("special", sidx))

            # 若同一列也含主持人資訊，嘗試把主持人加入 special（title="主持人"）
            host_names = _extract_host_names(speaker_text)
            for hn in host_names:
                h_idx = _add_special(specials, special_times, "主持人", talk_idx if talk_idx > 0 else 0, duration_val, hn or None, time_rng)
                if h_idx is not None:
                    timeline.append(("special", h_idx))
            continue

        # 否則視為 talk（需要時間；若時間寫在主題欄也試著抓）
        if time_rng is None:
            time_rng = _extract_time_range(topic_text)
        if time_rng is None:
            # 沒時間：避免把主持人/備註當 talk
            host_names = _extract_host_names(speaker_text)
            for hn in host_names:
                h_idx = _add_special(specials, special_times, "主持人", talk_idx if talk_idx > 0 else 0, None, hn or None, None)
                if h_idx is not None:
                    timeline.append(("special", h_idx))
            continue

        if first_time_start is None:
            first_time_start = time_rng[0]
        last_time_end = time_rng[1]

        talk_idx += 1
        tmin = _minutes_between(time_rng[0], time_rng[1])
        talk_durations.append(tmin)
        # 簡化講者名稱取第一行（通常是姓名）
        name_short = (speaker_text.split("\n")[0] if speaker_text else "").strip()
        talks.append({
            "no": talk_idx,  # 1-based during parsing
            "topic": topic_text if topic_text else "(未命名主題)",
            "name": name_short,
            "start_time": time_rng[0],
            "end_time": time_rng[1],
        })
        timeline.append(("talk", talk_idx))

        # 即便是 talk 列，也要額外檢查 speaker 欄是否含主持人，若有則加入 special (主持人)
        host_names = _extract_host_names(speaker_text)
        for hn in host_names:
            h_idx = _add_special(specials, special_times, "主持人", talk_idx, tmin, hn or None, time_rng)
            if h_idx is not None:
                timeline.append(("special", h_idx))

    # 只在「最後一列是 special 且像真正收尾」才把 after_speaker 設 999
    if timeline and timeline[-1][0] == "special":
        sidx = timeline[-1][1]
        title = specials[sidx]["title"]
        tr = special_times[sidx]
        end_matches_day_end = (tr is not None and last_time_end is not None and tr[1] == last_time_end)
        if end_matches_day_end or (title in END_TITLES):
            specials[sidx]["after_speaker"] = 999

    # 把 special_times 回填進 specials 裡面，方便後續輸出
    for idx, sp in enumerate(specials):
        tr = special_times[idx] if idx < len(special_times) else None
        if tr:
            sp["start_time"] = tr[0]
            sp["end_time"] = tr[1]
        else:
            sp["start_time"] = None
            sp["end_time"] = None

    # 講者時長取眾數
    speaker_minutes = 50
    if talk_durations:
        speaker_minutes = Counter(talk_durations).most_common(1)[0][0]

    if first_time_start is None:
        first_time_start = "09:00"
    if last_time_end is None:
        last_time_end = "17:00"

    agenda_settings = {
        "start_time": first_time_start,
        "end_time": last_time_end,
        "speaker_minutes": speaker_minutes,
        "special_sessions": specials,
    }

    # === Build speakers array combining host, talks, and time'd specials ===
    # Start from the ordered talks (講者) as the base list
    base_list: List[Dict[str, Any]] = []
    for t in talks:
        base_list.append({
            "type": "講者",
            "topic": t.get("topic"),
            "name": t.get("name"),
            "start_time": t.get("start_time"),
            "end_time": t.get("end_time")
        })

    # Helper to build a speaker-like dict from special
    def special_to_item(sp: Dict[str, Any]) -> Dict[str, Any]:
        title = sp.get("title")
        typ = title
        if title == "致詞":
            typ = "致詞人"
        elif title == "綜合討論":
            typ = "綜合討論"
        elif title in ("午間休息", "課間休息"):
            typ = "休息"
        elif title == "大合照":
            typ = "大合照"
        elif title == "主持人":
            typ = "主持人"
        # name handling
        sp_name_raw = sp.get("speaker") or ""
        sp_name = sp_name_raw.split("\n")[0].strip() if sp_name_raw else ""
        if not sp_name and typ == "綜合討論":
            sp_name = "所有講者"
        return {
            "type": typ,
            "topic": sp.get("title"),
            "name": sp_name,
            "start_time": sp.get("start_time"),
            "end_time": sp.get("end_time")
        }

    # Insert specials into base_list according to after_speaker when possible
    # after_speaker semantics: insert AFTER that talk index (talk numbering started at 1),
    # we map after_speaker -> insertion index in base_list (0-based).
    for sp in specials:
        item = special_to_item(sp)

        # Determine insert index
        after_sp = sp.get("after_speaker")
        insert_idx: int
        if isinstance(after_sp, int):
            if after_sp >= 999:
                insert_idx = len(base_list)  # append at end
            else:
                # after_sp is talk number (1-based) or 0 meaning before first talk
                # insertion index = after_sp (0-based)
                insert_idx = max(0, min(len(base_list), after_sp))
        else:
            # fallback: place by time if possible, else append
            st = sp.get("start_time")
            if st:
                found = False
                for i, t in enumerate(base_list):
                    tt = t.get("start_time")
                    if tt and tt > st:
                        insert_idx = i
                        found = True
                        break
                if not found:
                    insert_idx = len(base_list)
            else:
                insert_idx = len(base_list)

        # Avoid inserting empty-name主持人 duplicates: if item is 主持人 and name empty, put front only if no other host
        if item.get("type") == "主持人" and not item.get("name"):
            # check whether base_list already has a 主持人 with name
            existing_host = any((x.get("type") == "主持人" and x.get("name")) for x in base_list)
            if existing_host:
                # skip inserting duplicate anonymous host
                continue

        # Insert the item
        base_list.insert(insert_idx, item)

    # Ensure the first 主持人 (if any) is at index 0
    for i, it in enumerate(base_list):
        if it.get("type") == "主持人":
            if i != 0:
                base_list.insert(0, base_list.pop(i))
            break

    # Re-number and finalize speakers_out, setting 'no' and ensuring fields match expected keys
    speakers_out: List[Dict[str, Any]] = []
    for idx, it in enumerate(base_list):
        speakers_out.append({
            "no": idx,
            "type": it.get("type"),
            "topic": it.get("topic") or it.get("type"),
            "name": it.get("name") or "",
            "start_time": it.get("start_time"),
            "end_time": it.get("end_time")
        })

    # Build resulting minimal program structure (eventNames + agenda_settings + speakers)
    program_core = {
        "eventNames": [event_name],
        "agenda_settings": agenda_settings,
        "speakers": speakers_out,
        "activities_contacts": [{
            "ID": "",
            "id_number": "",
            "name": "",
            "gender": "",
            "title": "",
            "unity": [],
            "email": "",
            "contact_person": {"name": "", "email": ""},
            "group_by_program": [],
            "name_mail_use": "",
        }],
    }
    return program_core

def merge_with_metadata(parsed_program: Dict[str, Any],
                        *,
                        id: Optional[int] = None,
                        plan_name: Optional[str] = None,
                        date: Optional[str] = None,
                        locations: Optional[List[str]] = None,
                        instructors: Optional[List[str]] = None,
                        organizers: Optional[List[str]] = None,
                        co_organizers: Optional[List[str]] = None,
                        joint_organizers: Optional[List[str]] = None,
                        attachments: Optional[List[str]] = None,
                        registration_urls: Optional[List[str]] = None,
                        max_capacity: Optional[int] = None) -> Dict[str, Any]:
    """把 parse 出來的 program_core 與使用者提供的 metadata 合併成你想要的 JSON schema"""
    out: Dict[str, Any] = {}
    if id is not None:
        out["id"] = id
    if plan_name:
        out["planName"] = plan_name
    # 拷貝 eventNames + agenda + speakers
    out.update(parsed_program)
    # 加覆蓋或補充欄位
    if date:
        out["date"] = date
    if locations is not None:
        out["locations"] = locations
    if instructors is not None:
        out["instructors"] = instructors
    if organizers is not None:
        out["organizers"] = organizers
    if co_organizers is not None:
        out["coOrganizers"] = co_organizers
    if joint_organizers is not None:
        out["jointOrganizers"] = joint_organizers
    out["attachments"] = attachments or []
    out["registration_urls"] = registration_urls or []
    if max_capacity is not None:
        out["max_capacity"] = max_capacity
    # 若沒有某些欄位，保證欄位存在（方便下游系統）
    for k in ("planName","instructors","organizers","coOrganizers","jointOrganizers","attachments","registration_urls","max_capacity"):
        if k not in out:
            if k in ("attachments","registration_urls"):
                out[k] = []
            else:
                out[k] = []
    # activities_contacts 若不存在就補空
    out.setdefault("activities_contacts", parsed_program.get("activities_contacts", []))
    return out

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Parse a 3-column agenda DOCX (時間/主題/講者) into program JSON with speakers. 會輸出單一 JSON 物件 (not list).")
    ap.add_argument("--docx", required=True, help="Path to the agenda .docx")
    ap.add_argument("--event-name", required=True, help="Event name to put into eventNames")
    ap.add_argument("--out", required=True, help="Output program JSON file path")
    # optional metadata arguments to fill fields shown in your example
    ap.add_argument("--id", type=int, help="id field (optional)")
    ap.add_argument("--plan-name", help="planName (optional)")
    ap.add_argument("--date", help="date (YYYY-MM-DD)")
    ap.add_argument("--locations", help="comma-separated locations (or use '||' to separate)")
    ap.add_argument("--instructors", help="comma-separated instructors")
    ap.add_argument("--organizers", help="comma-separated organizers")
    ap.add_argument("--registration-urls", help="comma-separated registration URLs")
    ap.add_argument("--max-capacity", type=int, help="max capacity")
    args = ap.parse_args()

    parsed = parse_agenda(Path(args.docx), args.event_name)

    # locations splitting: support '||' (explicit separator) or comma
    if args.locations:
        if "||" in args.locations:
            locations = [s.strip() for s in args.locations.split("||") if s.strip()]
        else:
            locations = [s.strip() for s in args.locations.split(",") if s.strip()]
    else:
        locations = None

    instructors = [s.strip() for s in args.instructors.split(",")] if args.instructors else None
    organizers = [s.strip() for s in args.organizers.split(",")] if args.organizers else None
    reg_urls = [s.strip() for s in args.registration_urls.split(",")] if args.registration_urls else None

    merged = merge_with_metadata(parsed,
                                 id=args.id,
                                 plan_name=args.plan_name,
                                 date=args.date,
                                 locations=locations,
                                 instructors=instructors,
                                 organizers=organizers,
                                 registration_urls=reg_urls,
                                 max_capacity=args.max_capacity)

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
    print("[OK] Wrote {}".format(out_path))
