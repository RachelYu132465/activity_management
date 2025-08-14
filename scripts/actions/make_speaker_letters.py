# scripts/actions/make_speaker_letters.py
# pip install python-docx
# 放在檔案最前面

from __future__ import annotations

from pathlib import Path
import sys
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional

from docx import Document

# ---- 專案路徑（優先用你的 bootstrap，失敗則用後備猜測） ----
try:
    from scripts.core.bootstrap import initialize, OUTPUT_DIR, TEMPLATE_DIR
except Exception:
    BASE_DIR = Path(__file__).resolve().parents[2]
    OUTPUT_DIR = BASE_DIR / "output"
    TEMPLATE_DIR = BASE_DIR / "templates"
    def initialize(): OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

from scripts.core.build_mapping import get_event_speaker_mappings


# ---- 模板尋找（支援子資料夾）----
def find_template_file(template_filename: str) -> Path:
    p = TEMPLATE_DIR / template_filename
    if p.exists(): return p
    matches = list(TEMPLATE_DIR.rglob(template_filename))
    if matches: return matches[0]
    raise SystemExit(f"找不到模板：{template_filename}（已搜尋 {TEMPLATE_DIR} 子資料夾）")

# ---- DOCX 置換 ----
REPLACERS = [
    ("{{name}}", "name"),
    ("{{topic}}", "topic"),
    ("{{start_time}}", "start_time"),
    ("{{end_time}}", "end_time"),
    ("{{date}}", "date"),
    ("{{location_main}}", "location_main"),
    ("{{location_addr}}", "location_addr"),
    ("{{organization}}", "organization"),
    ("{{title}}", "title"),
    # 相容你的舊模板標記
    ("{{ activities_data.speakers. name}}", "name"),
    ("{{ activities_data.speakers. topic}}", "topic"),
    ("{{ activities_data.speakers. starttime}}", "start_time"),
    ("{{ activities_data.speakers. endtime}}", "end_time"),
    ("{{ program_data.. date}}", "date"),
    ("{{locations[0] }}", "location_main"),
    ("{{locations[1] }}", "location_addr"),
]

def render_docx(template_path: Path, out_path: Path, mapping: Dict[str, str]) -> None:
    doc = Document(str(template_path))

    def _apply(text: str) -> str:
        s = text
        for pat, key in REPLACERS:
            s = s.replace(pat, str(mapping.get(key, "")))
        return s

    # 段落
    for p in doc.paragraphs:
        if "{{" in p.text and "}}" in p.text:
            new = _apply(p.text)
            if new != p.text:
                for r in p.runs: r.text = ""
                p.add_run(new)

    # 表格
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text and "}}" in p.text:
                        new = _apply(p.text)
                        if new != p.text:
                            for r in p.runs: r.text = ""
                            p.add_run(new)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

# ---- 主流程 ----
def make_letters(event_name: str, template_filename: str,
                 out_dir: Optional[Path]=None,
                 filter_speaker_no: Optional[int]=None,
                 filter_speaker_name: Optional[str]=None) -> List[Path]:
    initialize()
    mappings = get_event_speaker_mappings(event_name)

    if filter_speaker_no is not None:
        mappings = [m for m in mappings if int(m.get("no", -1)) == int(filter_speaker_no)]
    if filter_speaker_name:
        target = filter_speaker_name.strip()
        mappings = [m for m in mappings if (m.get("name") or "").strip() == target]

    template_path = find_template_file(template_filename)
    out_base = out_dir or (OUTPUT_DIR / "letters")
    results: List[Path] = []

    for m in mappings:
        no = int(m.get("no", 0))
        pos = m.get("current_position") or {}
        mapping = {
            "name": m.get("name", ""),
            "topic": m.get("topic", ""),
            "start_time": m.get("start_time", ""),
            "end_time": m.get("end_time", ""),
            "date": m.get("date", ""),
            "location_main": m.get("location_main", ""),
            "location_addr": m.get("location_addr", ""),
            "organization": pos.get("organization") or m.get("organization", ""),
            "title": pos.get("title") or m.get("title", ""),
        }

        safe_name = m.get("safe_filename") or (m.get("name") or "TBD")
        out_name = f"{no:02d}_{safe_name}_敬請協助提供CV與簡報.docx"
        out_path = out_base / out_name
        render_docx(template_path, out_path, mapping)
        results.append(out_path)

    return results

if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="依 eventName 產出每位講者的《敬請協助提供CV與簡報》信件（Word）")
    ap.add_argument("--event", required=True, help="event name（與 program/activities 的 eventNames 任一相符即可）")
    ap.add_argument("--template", default="敬請協助提供CV與簡報.docx", help="模板檔名（放在 templates/ 或其子資料夾）")
    ap.add_argument("--outdir", default=None, help="輸出資料夾（預設 output/letters）")
    ap.add_argument("--speaker-no", type=int, default=None, help="（選配）只產出此講者編號")
    ap.add_argument("--speaker-name", default=None, help="（選配）只產出此講者姓名")
    args = ap.parse_args()

    outdir = Path(args.outdir) if args.outdir else None
    files = make_letters(args.event, args.template, outdir, args.speaker_no, args.speaker_name)
    print("=== DONE ===")
    for p in files:
        try:
            print("-", p.relative_to(Path(__file__).resolve().parents[2]))
        except Exception:
            print("-", p)
