# scripts/actions/generate_agenda_docx.py
# pip install python-docx
from __future__ import annotations
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import json
from datetime import datetime, timedelta
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from scripts.core.bootstrap import (
    initialize, DATA_DIR, OUTPUT_DIR
)

initialize()

# ---------- data loading ----------
def load_programs() -> List[Dict[str, Any]]:
    p = DATA_DIR / "shared" / "program_data.json"
    return json.loads(p.read_text(encoding="utf-8"))

def pick_event(programs: List[Dict[str, Any]], event_name: str) -> Dict[str, Any]:
    for prog in programs:
        names = prog.get("eventNames") or []
        if any(event_name == n for n in names):
            return prog
    raise SystemExit("找不到 event：{}".format(event_name))

# ---------- time helpers ----------

def gen_agenda_rows(event: Dict[str, Any]) -> List[Dict[str, str]]:
    """Generate agenda rows based on predefined speaker times.

    This version reads ``start_time`` and ``end_time`` directly from
    ``event['speakers']`` and inserts any special sessions defined in
    ``event['agenda_settings']['special_sessions']`` before or after the
    corresponding speakers.
    """

    def parse(t: str) -> datetime:
        return datetime.strptime(t, "%H:%M")

    def fmt(dt: datetime) -> str:
        return dt.strftime("%H:%M")

    specials = (event.get("agenda_settings", {}) or {}).get("special_sessions", []) or []
    speakers = event.get("speakers", []) or []

    rows: List[Dict[str, str]] = []

    # Specials before the first speaker
    first_sp = next((s for s in speakers if s.get("start_time")), None)
    if first_sp:
        first_start = parse(first_sp["start_time"])
        for s in specials:
            if int(s.get("after_speaker", -1)) == 0:
                # place special immediately before the first speaker
                start = first_start - timedelta(minutes=int(s.get("duration", 0)))
                rows.append({
                    "kind": "special",
                    "time": f"{fmt(start)}-{fmt(first_start)}",
                    "title": s.get("title", ""),
                    "speaker": "",
                })

    # Speakers and following specials
    for sp in speakers:
        start = sp.get("start_time")
        end = sp.get("end_time")
        if not start or not end:
            continue
        rows.append({
            "kind": "talk",
            "time": f"{start}-{end}",
            "title": sp.get("topic", ""),
            "speaker": sp.get("name", ""),
        })
        for s in specials:
            if int(s.get("after_speaker", -1)) == int(sp.get("no", -1)):
                start_dt = parse(end)
                end_dt = start_dt + timedelta(minutes=int(s.get("duration", 0)))
                rows.append({
                    "kind": "special",
                    "time": f"{fmt(start_dt)}-{fmt(end_dt)}",
                    "title": s.get("title", ""),
                    "speaker": "",
                })

    # Specials after the last speaker
    last_end = None
    for sp in reversed(speakers):
        if sp.get("end_time"):
            last_end = parse(sp["end_time"])
            break
    if last_end:
        for s in specials:
            if int(s.get("after_speaker", -1)) == 999:
                end_dt = last_end + timedelta(minutes=int(s.get("duration", 0)))
                rows.append({
                    "kind": "special",
                    "time": f"{fmt(last_end)}-{fmt(end_dt)}",
                    "title": s.get("title", ""),
                    "speaker": "",
                })
                last_end = end_dt

    return rows

# ---------- docx helpers ----------
# 顏色常數：tuple 格式 (R, G, B)
GREEN: Tuple[int, int, int] = (0x12, 0x6E, 0x2E)   # 深綠
WHITE: Tuple[int, int, int] = (0xFF, 0xFF, 0xFF)

def ensure_page_setup(doc: Document):
    """A4、窄邊界，盡量保證單頁"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

def set_cell_shading(cell, fill_rgb_tuple: Tuple[int, int, int]):
    """設定表格儲存格底色（用 6 碼 hex）"""
    r, g, b = fill_rgb_tuple
    hexcolor = "{:02X}{:02X}{:02X}".format(r, g, b)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tc_pr.append(shd)

def set_cell_text(cell, text: str, bold: bool = False,
                  color: Tuple[int, int, int] | None = None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, size_pt: float = 10.5):
    """設定表格儲存格文字屬性"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)  # tuple → RGBColor
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def add_agenda_table(doc: Document, rows: List[Dict[str, str]], title: str | None = None):
    # 標題（可選）
    if title:
        h = doc.add_paragraph(title)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.runs[0] if h.runs else h.add_run()
        run.font.size = Pt(14)
        run.bold = True

    # 建表（兩欄：時間、內容）
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"   # ← 這行
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 欄寬：A4 可用寬約 18cm（左右各 1.5cm 邊界）
    time_w = Cm(3.6)
    body_w = Cm(18.0 - 3.6)  # ≈14.4cm
    table.columns[0].width = time_w
    table.columns[1].width = body_w

    for r in rows:
        tr = table.add_row().cells
        # 時間欄
        set_cell_text(tr[0], r["time"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)

        if r["kind"] == "special":
            # 綠底白字，置中
            set_cell_shading(tr[0], GREEN)
            set_cell_shading(tr[1], GREEN)
            set_cell_text(tr[0], r["time"], bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(tr[1], r["title"], bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            # 內容：主題 + 換行 + 講者
            topic = r["title"].strip()
            name  = r["speaker"].strip()
            content = topic if not name else "{}\n{}".format(topic, name)
            set_cell_text(tr[1], content, align=WD_ALIGN_PARAGRAPH.LEFT)

    # 邊框（細線）
    tbl_pr = table._tbl.get_or_add_tblPr()
    tbl_borders = OxmlElement('w:tblBorders')
    for tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:{}'.format(tag))
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')       # 0.5pt
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

def export_agenda_docx(event: Dict[str, Any], out_path: Path):
    rows = gen_agenda_rows(event)
    doc = Document()
    ensure_page_setup(doc)

    # 標題可用第一個 eventName
    title = (event.get("eventNames") or ["議程"])[0]
    add_agenda_table(doc, rows, title=title)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print("[OK] 已輸出：{}".format(out_path))

# ---------- main ----------
if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="依 eventName 讀取 program_data.json，輸出一頁 A4 的議程表 DOCX")
    ap.add_argument("--event", required=True, help="event name（須與 program_data.json 的 eventNames 完全相同）")
    ap.add_argument("--out", default=str(OUTPUT_DIR / "letters" / "agenda.docx"), help="輸出路徑")
    args = ap.parse_args()

    programs = load_programs()
    event = pick_event(programs, args.event)
    export_agenda_docx(event, Path(args.out))
