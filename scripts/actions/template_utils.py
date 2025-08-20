# scripts/actions/template_utils.py
from __future__ import annotations
import re
import unicodedata
import logging
from datetime import datetime, date
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional

# Highlight 目前使用 WD_COLOR_INDEX.YELLOW（python-docx 的內建螢光黃）。如果要改色，可以改那裡的常數（但 python-docx 提供的顏色選項有限）。
#
# Jinja2 的 filter 最靈活；建議在環境有安裝 jinja2 時使用 {{ ... | cn_date | highlight }}。若沒有 jinja2，fallback 也支援 |。
#
# 如果你的模板裡 program_data.date 是像 "2025-09-22" 的字串，cn_date 能解析常見格式（ISO / %Y-%m-%d / %Y/%m/%d 等）。若格式特殊，請在 data source 提供 ISO 格式或 datetime.date。
#
# 如果 highlight 沒生效（例如使用的 python-docx 版本不支援 highlight_color），會安靜地忽略該屬性但文字仍會被放置。
# optional deps
try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
except ModuleNotFoundError:
    Document = None
    WD_COLOR_INDEX = None

try:
    import jinja2
except ModuleNotFoundError:
    jinja2 = None

# debug logging for template utils
logging.basicConfig(level=logging.DEBUG, format="%(levelname)s: %(message)s")

# tokens used internally to mark highlight ranges after rendering
_HL_START = "__<<HL>>__"
_HL_END = "__<</HL>>__"

# Email regex used by find_email_in_record
EMAIL_REGEX = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", re.UNICODE)


def _clean_cell_value(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, float):
        if val.is_integer():
            s = str(int(val))
        else:
            s = format(val, "f").rstrip("0").rstrip(".")
    else:
        s = str(val)
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\u3000", " ")
    s = re.sub(r"[\u200B-\u200F\uFEFF]", "", s)
    s = re.sub(r"[\x00-\x1F\x7F]", "", s)
    return s.strip()


def find_email_in_record(record: Dict[str, Any]) -> Optional[str]:
    """
    Robustly find an email address in a record's fields.
    """
    # 1) scan keys that look mail-like
    for k in list(record.keys()):
        if k is None:
            continue
        k_norm = _clean_cell_value(k).lower()
        try:
            if re.search(r"(mail|email|e-?mail|信箱|電子郵)", k_norm):
                raw = record.get(k)
                s = _clean_cell_value(raw)
                if not s:
                    continue
                parts = re.split(r"[;,/|\s]+", s)
                for p in parts:
                    p = p.strip().strip("()[]<>\"'")
                    if not p:
                        continue
                    if EMAIL_REGEX.fullmatch(p):
                        return p
                    m = EMAIL_REGEX.search(p)
                    if m:
                        return m.group(0)
                m = EMAIL_REGEX.search(s)
                if m:
                    return m.group(0)
        except Exception:
            continue

    # 2) fallback: scan all values
    for v in record.values():
        try:
            s = _clean_cell_value(v)
            if not s:
                continue
            parts = re.split(r"[;,/|\s]+", s)
            for p in parts:
                p = p.strip().strip("()[]<>\"'")
                if not p:
                    continue
                if EMAIL_REGEX.fullmatch(p):
                    return p
                m = EMAIL_REGEX.search(p)
                if m:
                    return m.group(0)
            m = EMAIL_REGEX.search(s)
            if m:
                return m.group(0)
        except Exception:
            continue

    return None


def sanitize_filename(s: str, max_len: int = 200) -> str:
    s = re.sub(r"[\\/:\*\?\"<>\|]+", "-", s or "")
    s = re.sub(r"\s+", " ", s).strip()
    return s[:max_len]


def find_template_file(template_filename: str, template_dir: Optional[Path] = None) -> Path:
    """
    Find template file under template_dir (or default templates/).
    """
    if template_dir is None:
        BASE_DIR = Path(__file__).resolve().parents[2]
        template_dir = BASE_DIR / "templates"
    p = Path(template_dir) / template_filename
    if p.exists():
        return p
    matches = list(Path(template_dir).rglob(template_filename))
    if matches:
        return matches[0]
    raise FileNotFoundError(f"找不到模板：{template_filename}（已搜尋 {template_dir} 及子資料夾）")


# -------------------- helpers: date formatting & highlight wrapper --------------------
def format_chinese_date(value: Any) -> str:
    """
    Accepts date/datetime or string. Returns "YYYY年M月D日(星期X)" with Chinese weekday.
    """
    if value is None:
        return ""
    # If already date/datetime
    dt = None
    if isinstance(value, datetime):
        dt = value.date()
    elif isinstance(value, date):
        dt = value
    else:
        s = str(value).strip()
        # Try ISO
        try:
            dt = datetime.fromisoformat(s).date()
        except Exception:
            # try common formats
            fmts = ["%Y-%m-%d", "%Y/%m/%d", "%Y%m%d", "%Y.%m.%d", "%Y %m %d"]
            for f in fmts:
                try:
                    dt = datetime.strptime(s, f).date()
                    break
                except Exception:
                    continue
    if dt is None:
        # fallback: return original string
        return str(value)
    # weekday: Monday=0 -> 一 ... Sunday=6 -> 日
    wmap = ["一", "二", "三", "四", "五", "六", "日"]
    weekday_char = wmap[dt.weekday()] if 0 <= dt.weekday() <= 6 else ""
    return f"{dt.year}年{dt.month}月{dt.day}日(星期{weekday_char})"


def _wrap_highlight(s: str) -> str:
    """Wrap s with internal highlight markers."""
    return f"{_HL_START}{s}{_HL_END}"


# -------------------- core: render docx template / body --------------------
def render_docx_template(template_path: Path, out_path: Path, mapping: Dict[str, Any],
                         replacers: Optional[List[Tuple[str, str]]] = None) -> None:
    """
    Render a .docx template by performing paragraph-level replacements.
    Supports special pipe syntax in placeholders, e.g. {{ program_data.date|cn_date|highlight }}.
    If 'highlight' is applied, the output .docx will show that text highlighted (yellow).
    """
    if Document is None:
        raise ModuleNotFoundError("python-docx is required for render_docx_template")
    doc = Document(str(template_path))

    # simple apply_text that supports pipes
    def apply_text(text: str) -> str:
        if not text:
            return text
        out = text
        if replacers:
            for pat, key in replacers:
                out = out.replace(pat, str(mapping.get(key, "")))
            # convert literal \n to newline
            return out.replace("\\n", "\n")

        # placeholder pattern: {{ ... }}
        placeholder_re = re.compile(r"\{\{\s*(.*?)\s*\}\}")

        def _resolve(expr: str) -> str:
            # support piping: key|filter1|filter2
            parts = [p.strip() for p in expr.split("|") if p.strip()]
            if not parts:
                return ""
            key_expr = parts[0]
            # Try dotted/indexed resolution
            val = _resolve_path(mapping, key_expr)
            if val is None:
                val = mapping.get(key_expr)
            s = "" if val is None else str(val)
            # apply filters
            for f in parts[1:]:
                if f in ("cn_date", "cnDate", "format_date"):
                    s = format_chinese_date(s)
                elif f in ("hl", "highlight"):
                    s = _wrap_highlight(s)
                else:
                    # unknown filter: ignore
                    pass
            return s

        return placeholder_re.sub(lambda m: _resolve(m.group(1)), out).replace("\\n", "\n")

    # Helper: write paragraph text but respect highlight markers
    def _write_para_with_highlight(para, rendered_text: str):
        # split into segments [(text, highlighted_bool), ...]
        segs: List[Tuple[str, bool]] = []
        t = rendered_text
        idx = 0
        while True:
            i = t.find(_HL_START, idx)
            if i == -1:
                # rest is normal
                segs.append((t[idx:], False))
                break
            # prefix normal
            if i > idx:
                segs.append((t[idx:i], False))
            j = t.find(_HL_END, i + len(_HL_START))
            if j == -1:
                # no matching end: treat remainder as normal
                segs.append((t[i:], False))
                break
            inner = t[i + len(_HL_START): j]
            segs.append((inner, True))
            idx = j + len(_HL_END)
            if idx >= len(t):
                break
        # clear existing runs' text
        if para.runs:
            for r in para.runs:
                r.text = ""
            base_run = para.runs[0]
        else:
            base_run = para.add_run("")
        # write segments: first one goes to base_run, others as new runs
        for i, (txt, hl) in enumerate(segs):
            if txt == "":
                continue
            if i == 0:
                run = base_run
                run.text = txt
            else:
                run = para.add_run(txt)
            # apply highlight if requested and python-docx supports it
            if hl and WD_COLOR_INDEX is not None:
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    # ignore if highlight not supported
                    pass

    # paragraphs
    for para in doc.paragraphs:
        if "{{" in para.text or "}}" in para.text:
            new = apply_text(para.text)
            _write_para_with_highlight(para, new)

    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text or "}}" in para.text:
                        new = apply_text(para.text)
                        _write_para_with_highlight(para, new)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# Resolve dotted/indexed path like "program_data.locations[0]"
_key_token_re = re.compile(r'([A-Za-z0-9_]+)|\[(\d+)\]')


def _resolve_path(mapping: Any, expr: str) -> Optional[Any]:
    cur = mapping
    for m in _key_token_re.finditer(expr):
        name, idx = m.group(1), m.group(2)
        if name:
            try:
                if isinstance(cur, dict):
                    cur = cur.get(name)
                else:
                    cur = getattr(cur, name, None)
            except Exception:
                return None
        else:
            try:
                cur = cur[int(idx)]
            except Exception:
                return None
        if cur is None:
            return None
    return cur


def _apply_filters_to_value(value: Any, filters: List[str]) -> str:
    s = "" if value is None else str(value)
    for f in filters:
        if f in ("cn_date", "cnDate", "format_date"):
            s = format_chinese_date(s)
        elif f in ("hl", "highlight"):
            s = _wrap_highlight(s)
        else:
            # unknown filter: ignore
            pass
    return s


def render_body_from_template(template_path: Path, context: Dict[str, Any]) -> str:
    """
    Render textual body from a docx template. Returns string (paragraphs joined with \n).
    Uses Jinja2 if available (registers 'cn_date' and 'highlight' filters), else falls back
    to a manual paragraph-level replacement that also understands pipe filters:
      {{ program_data.date | cn_date | highlight }}
    For highlight, the returned docx will contain highlighted runs; this function returns
    plain-text body with highlighted markers removed (but keep line breaks).
    """
    if Document is None:
        raise ModuleNotFoundError("python-docx is required for template rendering")
    doc = Document(str(template_path))

    # Setup jinja2 env and filters if available
    jenv = None
    if jinja2 is not None:
        try:
            jenv = jinja2.Environment(undefined=jinja2.StrictUndefined)
            # register filters
            jenv.filters["cn_date"] = format_chinese_date
            # highlight filter wraps with internal markers; actual highlight applied when writing runs
            jenv.filters["highlight"] = lambda s: _wrap_highlight("" if s is None else str(s))
            jenv.filters["hl"] = jenv.filters["highlight"]
        except Exception:
            jenv = None

    placeholder_re = re.compile(r"\{\{\s*(.*?)\s*\}\}")

    def render_text(text: str) -> str:
        if not text:
            return text
        # try jinja2 first (paragraph-level)
        if jenv is not None:
            try:
                tmpl = jenv.from_string(text)
                return tmpl.render(**context)
            except Exception:
                # fall back to manual replacement below
                pass

        # manual replacement with pipe support
        def _repl(m: re.Match) -> str:
            expr = m.group(1).strip()
            parts = [p.strip() for p in expr.split("|") if p.strip()]
            if not parts:
                return m.group(0)
            key_expr = parts[0]
            filters = parts[1:]
            # resolve key
            val = _resolve_path(context, key_expr)
            if val is None:
                # fallback to direct key
                val = context.get(key_expr)
            return _apply_filters_to_value(val, filters)

        return placeholder_re.sub(_repl, text)

    # When writing back to docx we must honour highlight markers and replace them with runs with highlight.
    def _write_para_with_highlight(para, rendered_text: str):
        # same logic as in render_docx_template
        segs: List[Tuple[str, bool]] = []
        t = rendered_text
        idx = 0
        while True:
            i = t.find(_HL_START, idx)
            if i == -1:
                segs.append((t[idx:], False))
                break
            if i > idx:
                segs.append((t[idx:i], False))
            j = t.find(_HL_END, i + len(_HL_START))
            if j == -1:
                segs.append((t[i:], False))
                break
            inner = t[i + len(_HL_START): j]
            segs.append((inner, True))
            idx = j + len(_HL_END)
            if idx >= len(t):
                break

        # clear runs
        if para.runs:
            for r in para.runs:
                r.text = ""
            base_run = para.runs[0]
        else:
            base_run = para.add_run("")
        for i, (txt, hl) in enumerate(segs):
            if txt == "":
                continue
            if i == 0:
                run = base_run
                run.text = txt
            else:
                run = para.add_run(txt)
            if hl and WD_COLOR_INDEX is not None:
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    pass

    # process paragraphs
    for para in doc.paragraphs:
        if "{{" in para.text or "{%" in para.text:
            try:
                new_text = render_text(para.text)
            except Exception:
                new_text = para.text
            # convert literal \n to real newline
            if new_text is not None:
                new_text = new_text.replace("\\n", "\n")
            _write_para_with_highlight(para, new_text)

    # tables / cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text or "{%" in para.text:
                        try:
                            new_text = render_text(para.text)
                        except Exception:
                            new_text = para.text
                        if new_text is not None:
                            new_text = new_text.replace("\\n", "\n")
                        _write_para_with_highlight(para, new_text)

    # Produce a plain-text body (paragraphs joined with real newline). Highlight markers removed.
    body_lines = [p.text for p in doc.paragraphs]
    return "\n".join(body_lines).strip()
