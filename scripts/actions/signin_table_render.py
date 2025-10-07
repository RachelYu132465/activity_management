"""Utilities for rendering speaker sign-in tables to Word documents.

The module exposes :func:`render_signin_table` for full control as well as the
shortcut :func:`render_signin_table_paginated` that other Python modules can
import when they simply need to supply a ``rows_per_page`` limit::

    from scripts.actions.signin_table_render import (
        SignInDocumentContext,
        SignInRow,
        render_signin_table_paginated,
    )

    context = SignInDocumentContext(plan_name="Plan", event_name="Event")
    rows = [SignInRow(name="Alice"), SignInRow(name="Bob")]
    render_signin_table_paginated(context, rows, "output.docx", rows_per_page=15)
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from typing import Iterable, Protocol, Sequence, Tuple, runtime_checkable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# Default layout constants (cm / pt)
TITLE_PT = 16
FONT_PT = 14
LEFT_RIGHT_MARGIN_CM = 1.5
COL0_FIXED_CM = 8.0
COL1_FIXED_CM = 8.0
HDR_HEIGHT_CM = 0.9
DATA_ROW_HEIGHT_CM = 3.5


@dataclass
class SignInRow:
    """Single row entry for the sign-in table."""

    topic: str = ""
    name: str = ""
    title: str = ""
    organization: str = ""


@dataclass
class SignInDocumentContext:
    """Document level metadata used when rendering the sign-in sheet."""

    plan_name: str = ""
    event_name: str = ""
    date_display: str | None = None
    subtitle: str = "講員簽到單"


@dataclass
class SignInRenderResult:
    """Return information after rendering a sign-in sheet."""

    output_path: Path
    page_width_cm: float
    available_width_cm: float
    columns_cm: Tuple[float, float, float]


def _set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    run.font.size = Pt(size_pt)
    run.bold = bold


def _set_table_cell_margins(
    table, left_cm: float = 0.1, right_cm: float = 0.1, top_pt: int = 0, bottom_pt: int = 0
):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tcMar = tblPr.find(qn("w:tblCellMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tblCellMar")
        tblPr.append(tcMar)

    def _set_node(name: str, cm_val: float, parent: OxmlElement):
        node = parent.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            parent.append(node)
        node.set(qn("w:w"), str(int(cm_val * 567)))
        node.set(qn("w:type"), "dxa")

    _set_node("left", left_cm, tcMar)
    _set_node("right", right_cm, tcMar)
    _set_node("top", (top_pt / 72.0) * 2.54 if top_pt else 0, tcMar)
    _set_node("bottom", (bottom_pt / 72.0) * 2.54 if bottom_pt else 0, tcMar)


def _set_table_total_width(table, total_cm: float):
    tbl = table._tbl
    tblPr = tbl.tblPr
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(total_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _safe_set_row_height(row, height_cm: float, preferred_rule: str = "EXACT"):
    row.height = Cm(height_cm)
    rule_value = None
    if hasattr(WD_ROW_HEIGHT_RULE, preferred_rule):
        rule_value = getattr(WD_ROW_HEIGHT_RULE, preferred_rule)
    else:
        for alt in ("EXACT", "AT_LEAST", "AUTO"):
            if hasattr(WD_ROW_HEIGHT_RULE, alt):
                rule_value = getattr(WD_ROW_HEIGHT_RULE, alt)
                break
    if rule_value is not None:
        try:
            row.height_rule = rule_value
        except Exception:
            pass


def _set_row_height_exact(row, height_cm: float):
    row.height = Cm(height_cm)
    rule_value = None
    if hasattr(WD_ROW_HEIGHT_RULE, "EXACT"):
        rule_value = WD_ROW_HEIGHT_RULE.EXACT
    elif hasattr(WD_ROW_HEIGHT_RULE, "AT_LEAST"):
        rule_value = WD_ROW_HEIGHT_RULE.AT_LEAST
    elif hasattr(WD_ROW_HEIGHT_RULE, "AUTO"):
        rule_value = WD_ROW_HEIGHT_RULE.AUTO

    if rule_value is not None:
        try:
            row.height_rule = rule_value
        except Exception:
            pass


def _set_table_rows_height(table, height_cm: float):
    for r in table.rows:
        _set_row_height_exact(r, height_cm)


def _set_row_height_auto(row):
    if hasattr(WD_ROW_HEIGHT_RULE, "AUTO"):
        try:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        except Exception:
            pass


def _set_cell_vertical_center(cell):
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        pass


def _set_cell_background(cell, color_hex: str):
    color = color_hex.lstrip("#")
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.append(tcPr)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color.upper())
    tcPr.append(shd)


def _set_run_color_black(run):
    try:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except Exception:
        pass


def _set_repeat_table_header(row) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        trPr.append(tbl_header)


def _sanitize_text(value: str | None) -> str:
    return (value or "").strip()


@runtime_checkable
class SupportsSignInRow(Protocol):
    """Protocol for objects that can provide a :class:`SignInRow`."""

    def to_signin_row(self) -> "SignInRow":
        """Return the :class:`SignInRow` representation of the object."""


def _coerce_signin_rows(entries: Iterable[SignInRow | SupportsSignInRow]) -> list[SignInRow]:
    rows: list[SignInRow] = []
    for entry in entries:
        if isinstance(entry, SignInRow):
            rows.append(entry)
        elif isinstance(entry, SupportsSignInRow):
            row = entry.to_signin_row()
            if not isinstance(row, SignInRow):
                raise TypeError(
                    "to_signin_row() must return a SignInRow instance, "
                    f"got {type(row)!r}"
                )
            rows.append(row)
        else:
            raise TypeError(
                "Unsupported entry type for sign-in rendering: "
                f"{type(entry)!r}. Provide SignInRow or SupportsSignInRow instances."
            )
    return rows


def render_signin_table(
    context: SignInDocumentContext,
    speakers: Sequence[SignInRow | SupportsSignInRow],

    output_path: Path | str,
    *,
    title_pt: int = TITLE_PT,
    font_pt: int = FONT_PT,
    left_right_margin_cm: float = LEFT_RIGHT_MARGIN_CM,
    col0_fixed_cm: float = COL0_FIXED_CM,
    col1_fixed_cm: float = COL1_FIXED_CM,
    data_row_height_cm: float = DATA_ROW_HEIGHT_CM,
    header_height_cm: float = HDR_HEIGHT_CM,
    auto_adjust_dimensions: bool = True,
    rows_per_page: int | None = None,
) -> SignInRenderResult:
    """Render a sign-in table document and save it to ``output_path``.

    Parameters
    ----------
    rows_per_page:
        Optional limit of data rows for each page. When provided, the
        renderer will insert a page break after the given number of rows and
        recreate the table with the header so every page starts with the
        header row. When ``None`` (default) all rows are rendered into a
        single table, matching the legacy behaviour.
    auto_adjust_dimensions:
        When ``True`` (default) column widths and row heights are left to
        Word's automatic layout engine so content determines the final size.
        When ``False`` the legacy fixed widths and heights defined by the
        module constants are applied.
    """

    normalized = _coerce_signin_rows(speakers)
    rows: list[SignInRow] = [
        row
        for row in normalized
        if any(_sanitize_text(getattr(row, field)) for field in ("name", "topic", "title", "organization"))
    ]

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    normal_font.size = Pt(font_pt)

    if context.plan_name:
        p_plan = doc.add_paragraph()
        p_plan.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_plan.add_run(context.plan_name)
        _set_run_font(run, title_pt, bold=True)

    if context.event_name:
        p_event = doc.add_paragraph()
        p_event.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_event.add_run(context.event_name)
        _set_run_font(run, title_pt, bold=True)

    if context.date_display:
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_date.add_run(f"({context.date_display})")
        _set_run_font(run, title_pt, bold=True)

    if context.subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(context.subtitle)
        _set_run_font(run, title_pt, bold=True)

    if doc.sections:
        doc.sections[0].left_margin = Cm(left_right_margin_cm)
        doc.sections[0].right_margin = Cm(left_right_margin_cm)

    page_width_cm = float(doc.sections[0].page_width) / float(Cm(1)) if doc.sections else 0.0
    available_cm = page_width_cm - (left_right_margin_cm * 2)

    col0 = float(col0_fixed_cm)
    col1 = float(col1_fixed_cm)
    col2 = available_cm - (col0 + col1)

    total_req = col0 + col1 + (col2 if col2 > 0 else 0)
    if available_cm > 0 and total_req > available_cm:
        scale = available_cm / total_req
        col0 *= scale
        col1 *= scale
        col2 = max(1.5, available_cm - (col0 + col1))

    if col2 < 1.5:
        col2 = 1.5
        if col0 + col1 + col2 > available_cm and available_cm > 0:
            remain = max(0, available_cm - col2)
            ratio = col0 / (col0 + col1) if (col0 + col1) else 0.6
            col0 = remain * ratio
            col1 = remain - col0

    rows_per_page = rows_per_page if rows_per_page and rows_per_page > 0 else None

    def _create_table():
        table = doc.add_table(rows=1, cols=3, style="Table Grid")
        table.autofit = bool(auto_adjust_dimensions)
        if auto_adjust_dimensions:
            try:
                table.allow_autofit = True
            except Exception:
                pass
        else:
            table.autofit = False
            table.columns[0].width = Cm(col0)
            table.columns[1].width = Cm(col1)
            table.columns[2].width = Cm(col2)
            _set_table_total_width(table, col0 + col1 + col2)
        _set_table_cell_margins(table, left_cm=0.12, right_cm=0.12)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.left_indent = Cm(0)
        except Exception:
            pass

        hdr_row = table.rows[0]
        if auto_adjust_dimensions:
            _set_row_height_auto(hdr_row)
        else:
            _safe_set_row_height(hdr_row, header_height_cm)
             _set_repeat_table_header(hdr_row)
            _set_row_height_exact(table.rows[0], header_height_cm)
       

        hdr_cells = hdr_row.cells
        headers = ["主題 Topic", "姓名 Name", "簽到 Sign-in"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            _set_run_font(run, font_pt, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_vertical_center(cell)
            _set_cell_background(cell, "#BFBFBF")

        return table

    def _populate_row(table_row, entry: SignInRow):
        if auto_adjust_dimensions:
            _set_row_height_auto(table_row)
        else:
            try:
                _set_row_height_exact(table_row, data_row_height_cm)
            except Exception:
                _safe_set_row_height(table_row, data_row_height_cm)


        row_cells = table_row.cells
        topic_val = _sanitize_text(entry.topic) or "\u00A0"
        c0 = row_cells[0]
        p0 = c0.paragraphs[0] if c0.paragraphs else c0.add_paragraph()
        for r in list(p0.runs):
            try:
                r._element.getparent().remove(r._element)
            except Exception:
                pass
        p0.add_run(topic_val)
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
        _set_cell_vertical_center(c0)

        name_val = _sanitize_text(entry.name) or "\u00A0"
        title_val = _sanitize_text(entry.title)
        org_val = _sanitize_text(entry.organization)

        c1 = row_cells[1]
        p1 = c1.paragraphs[0] if c1.paragraphs else c1.add_paragraph()
        for r in list(p1.runs):
            try:
                r._element.getparent().remove(r._element)
            except Exception:
                pass
        r_name = p1.add_run(name_val)
        _set_run_font(r_name, font_pt, bold=True)
        if title_val:
            r_title = p1.add_run(f" {title_val}")
            _set_run_font(r_title, font_pt, bold=False)
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
        except Exception:
            pass

        if org_val:
            org_p = c1.add_paragraph()
            org_p.text = org_val
            if org_p.runs:
                _set_run_font(org_p.runs[0], font_pt, bold=False)
            try:
                org_p.paragraph_format.space_before = Pt(0)
                org_p.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
        _set_cell_vertical_center(c1)

        c2 = row_cells[2]
        p2 = c2.paragraphs[0] if c2.paragraphs else c2.add_paragraph()
        for r in list(p2.runs):
            try:
                r._element.getparent().remove(r._element)
            except Exception:
                pass
        r2 = p2.add_run("\u00A0")
        try:
            r2.font.size = Pt(font_pt)
        except Exception:
            pass
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
        _set_cell_vertical_center(c2)

    chunks: list[list[SignInRow]]
    if rows_per_page is None:
        chunks = [rows]
    else:
        chunks = [rows[i : i + rows_per_page] for i in range(0, len(rows), rows_per_page)]

    for index, chunk in enumerate(chunks):
        if index > 0:
            doc.add_page_break()
        table = _create_table()
        for entry in chunk:
            table_row = table.add_row()
            _populate_row(table_row, entry)

    if context.date_display:
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(context.date_display)
        _set_run_font(run, font_pt, bold=False)

    doc.save(out_path)
    return SignInRenderResult(
        output_path=out_path,
        page_width_cm=page_width_cm,
        available_width_cm=available_cm,
        columns_cm=(col0, col1, col2),
    )


def render_signin_table_paginated(
    context: SignInDocumentContext,
    speakers: Sequence[SignInRow | SupportsSignInRow],
    output_path: Path | str,
    *,
    rows_per_page: int,
    **kwargs,
) -> SignInRenderResult:
    """Convenience helper that forwards to :func:`render_signin_table`.

    Other Python modules that only need to specify a ``rows_per_page`` limit
    can import and call this helper instead of remembering the keyword name::

        from scripts.actions.signin_table_render import render_signin_table_paginated

        render_signin_table_paginated(context, rows, "output.docx", rows_per_page=15)

    All other keyword arguments accepted by :func:`render_signin_table` may be
    supplied through ``**kwargs``.
    """

    return render_signin_table(
        context,
        speakers,
        output_path,
        rows_per_page=rows_per_page,
        **kwargs,
    )


__all__ = [
    "SignInRow",
    "SignInDocumentContext",
    "SignInRenderResult",
    "SupportsSignInRow",
    "render_signin_table",
    "render_signin_table_paginated",
    "TITLE_PT",
    "FONT_PT",
    "LEFT_RIGHT_MARGIN_CM",
    "COL0_FIXED_CM",
    "COL1_FIXED_CM",
    "HDR_HEIGHT_CM",
    "DATA_ROW_HEIGHT_CM",
]
