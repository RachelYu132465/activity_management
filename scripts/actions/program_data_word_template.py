"""
Render a Word (.docx) template by replacing placeholders like:
    {{ program_data.all_topics }}
    {{ program_data.course_summary }}
and more if you add them to the mapping.

Usage:
    python render_program_template.py ^
        --template templates\\program_template.docx ^
        --data data\\shared\\program_data.json ^
        --program-id 7 ^
        --output output\\program_rendered_id7.docx
"""

from __future__ import annotations
import argparse
import json
import re
from pathlib import Path
from typing import Any, Dict, Iterator, List, Sequence, Tuple

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("請先安裝 python-docx：pip install python-docx") from exc


# ------------------------- Topic extraction utils ------------------------------
_EXCLUDE_TOPICS = {
    "報到", "簽到", "開場", "引言", "致詞", "主持", "主持人", "合影", "茶敘", "交流",
    "綜合討論", "Q&A", "Q＆A", "QA", "午間休息", "課間休息", "中場休息",
    "閉幕", "結語", "頒發感謝狀", "賦歸",
}

def _is_valid_topic(item: Any) -> bool:
    if not isinstance(item, dict):
        return False
    topic = str(item.get("topic") or "").strip()
    if not topic or topic in _EXCLUDE_TOPICS:
        return False
    item_type = (item.get("type") or "").lower()
    if item_type in {"break", "rest", "host", "opening", "closing", "photo", "admin"}:
        return False
    return True

def extract_topics(program: dict[str, Any]) -> List[str]:
    speakers = program.get("speakers") or []
    out: List[str] = []
    seen: set[str] = set()
    for row in speakers:
        if _is_valid_topic(row):
            t = str(row["topic"]).strip()
            if t and t not in seen:
                seen.add(t)
                out.append(t)
    return out

def build_all_topics(program: dict[str, Any]) -> str:
    return "、".join(extract_topics(program)) or ""

def build_course_summary(program: dict[str, Any]) -> str:
    topics = extract_topics(program)
    if not topics:
        return "課程內容包括：〈尚無可列示之課程講題〉。"
    return f"課程內容包括：「{'」、「'.join(topics)}」。"


# ------------------------------- I/O helpers -----------------------------------
def load_programs(path: Path) -> Sequence[dict[str, Any]]:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise SystemExit(f"找不到 program_data：{path}") from exc
    if isinstance(payload, list):
        return payload
    if isinstance(payload, dict):
        return [payload]
    raise SystemExit("program_data.json 內容必須是 JSON 物件或陣列")

def select_program(programs: Sequence[dict[str, Any]], program_id: int | None) -> dict[str, Any]:
    if program_id is None:
        if not programs:
            raise SystemExit("program_data 為空")
        return programs[0]
    for p in programs:
        if p.get("id") == program_id:
            return p
    raise SystemExit(f"找不到指定 id 的 program：{program_id}")


# ------------------------------ Word helpers -----------------------------------
def _iter_paragraphs_in_table(table) -> Iterator[Any]:
    for row in table.rows:
        for cell in row.cells:
            # paragraphs
            for p in cell.paragraphs:
                yield p
            # nested tables
            for t2 in cell.tables:
                yield from _iter_paragraphs_in_table(t2)

def iter_all_paragraphs(doc: Document) -> Iterator[Any]:
    """Yield all paragraphs in body, tables, headers, footers."""
    # body paragraphs
    for p in doc.paragraphs:
        yield p
    # body tables
    for t in doc.tables:
        yield from _iter_paragraphs_in_table(t)
    # headers/footers
    for section in doc.sections:
        hdr = section.header
        if hdr:
            for p in hdr.paragraphs:
                yield p
            for t in hdr.tables:
                yield from _iter_paragraphs_in_table(t)
        ftr = section.footer
        if ftr:
            for p in ftr.paragraphs:
                yield p
            for t in ftr.tables:
                yield from _iter_paragraphs_in_table(t)

def replace_placeholders_in_paragraph(paragraph, patterns: List[Tuple[re.Pattern, str]]) -> bool:
    """
    Replace placeholders in one paragraph. Returns True if changed.
    NOTE: This clears runs and writes a single run for the whole paragraph.
          If you had mixed formatting inside the same paragraph (bold/italic on different parts),
          that formatting may be lost. Keep placeholders in their own runs/lines to avoid issues.
    """
    if not paragraph.runs:
        return False
    text = "".join(run.text for run in paragraph.runs)
    new_text = text
    for pat, repl in patterns:
        new_text = pat.sub(repl, new_text)
    if new_text != text:
        # clear runs
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph.runs[0].text = ""
            paragraph.runs[0].font.bold = paragraph.runs[0].font.bold  # no-op to keep style object
            paragraph.runs[0].element.getparent().remove(paragraph.runs[0].element)
        # write back as one run
        paragraph.add_run(new_text)
        return True
    return False

def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> int:
    """
    Replace placeholders across the document.
    Supports whitespace inside braces like {{  program_data.all_topics  }}.
    Returns number of replacements performed.
    """
    # Build compiled regex patterns for each key
    patterns: List[Tuple[re.Pattern, str]] = []
    for key, value in mapping.items():
        # Escape dots and build pattern allowing internal spaces: {{\s*key\s*}}
        # We anchor to full token to avoid partial hits.
        escaped = re.escape(key)
        pat = re.compile(r"\{\{\s*" + escaped + r"\s*\}\}")
        patterns.append((pat, value))

    count = 0
    for p in iter_all_paragraphs(doc):
        before = "".join(r.text for r in p.runs)
        changed = replace_placeholders_in_paragraph(p, patterns)
        if changed:
            after = "".join(r.text for r in p.runs)
            # rough count: sum of occurrences replaced for all patterns
            for pat, _ in patterns:
                count += len(pat.findall(before))
    return count


# ----------------------------------- CLI ---------------------------------------
def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Render a Word template by replacing {{ ... }} placeholders.")
    ap.add_argument("--template", type=Path, required=True, help="Path to the .docx template")
    ap.add_argument("--data", type=Path, required=True, help="Path to program_data.json")
    ap.add_argument("--program-id", type=int, help="Specific program id to render (default: first one)")
    ap.add_argument("--output", type=Path, required=True, help="Output .docx path")
    return ap.parse_args()

def main() -> None:
    args = parse_args()

    programs = load_programs(args.data)
    program = select_program(programs, args.program_id)

    # Build values for placeholders
    mapping: Dict[str, str] = {
        "program_data.all_topics": build_all_topics(program),
        "program_data.course_summary": build_course_summary(program),
        # 你也可以在這裡擴充更多欄位，例如：
        "program_data.planName": str(program.get("planName", "")),
        "program_data.eventNames[0]": str((program.get("eventNames") or [""])[0]),
        "program_data.date": str(program.get("date") or ""),
        "program_data.locations[0]": str((program.get("locations") or [""])[0]),
    }

    # Load template and replace
    doc = Document(str(args.template))
    replacements = replace_placeholders(doc, mapping)

    # Save
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"[OK] 產生完成：{args.output}（共替換 {replacements} 處）")

if __name__ == "__main__":
    main()
