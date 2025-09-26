#!/usr/bin/env python3
"""Render program handbook to a Word document.

This script loads ``program_data.json`` and ``influencer_data.json`` from the
``data/shared`` directory, selects a program by ``--program-id`` (defaults to the
first program) and creates a simple `.docx` file in the ``output`` directory.

The goal is to mirror ``templates/template.html`` but for Word output.  The
layout is intentionally simple so that the generated document remains
readable even without HTML rendering support.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import sys

from __init__ import format_date


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
# Project helpers
from scripts.core.bootstrap import DATA_DIR, OUTPUT_DIR, initialize
from scripts.actions.influencer import build_people
from docx.shared import Cm
from docx.shared import Pt, RGBColor
# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def load_program(program_id: int | None) -> Dict[str, Any]:
    """Return the program matching ``program_id`` from program_data.json.

    If ``program_id`` is ``None`` or not found, the first program entry is
    returned.
    """
    data_file = DATA_DIR / "shared" / "program_data.json"
    programs_raw = json.loads(data_file.read_text(encoding="utf-8"))

    # program_data.json may contain a list or a single dict
    if isinstance(programs_raw, list):
        if program_id is not None:
            for prog in programs_raw:
                try:
                    if int(prog.get("id", -1)) == program_id:
                        return prog
                except (TypeError, ValueError):
                    continue
        return programs_raw[0] if programs_raw else {}
    elif isinstance(programs_raw, dict):
        return programs_raw
    return {}


def build_schedule(event: Dict[str, Any]) -> List[Dict[str, str]]:
    """Build schedule rows from speaker information.

    The result is a list of dictionaries containing ``time``, ``topic`` and
    ``speaker`` keys, using each speaker's ``start_time`` and ``end_time``
    fields.  Special sessions are ignored so that only the regular speaker
    timetable is returned.
    """
    def time_range(start: str | None, end: str | None) -> str:
        if start and end:
            return f"{start}-{end}"
        return start or end or ""

    speakers = event.get("speakers", []) or []

    rows: List[Dict[str, str]] = []

    host = next((sp for sp in speakers if sp.get("type") == "主持人"), None)
    if host:
        text = " ".join(
            filter(
                None,
                [
                    time_range(host.get("start_time"), host.get("end_time")),
                    host.get("topic"),
                    host.get("name"),
                ],
            )
        )
        rows.append({"kind": "host", "time": "", "topic": text, "speaker": ""})

    for sp in speakers:
        if sp.get("type") == "主持人":
            continue
        start = sp.get("start_time")
        end = sp.get("end_time")
        rows.append(
            {
                "kind": "talk",
                "time": time_range(start, end),
                "topic": sp.get("topic", ""),
                "speaker": sp.get("name", ""),
                "no":sp.get("no", "")
            }
        )
    return rows


def set_run_font(run, size_pt: int, bold: bool = False) -> None:
    """Apply project font settings to ``run``.

    Chinese characters use Microsoft JhengHei while Latin characters use
    Times New Roman.  ``size_pt`` is the font size in points.
    """
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size_pt)
    run.bold = bold


def render_cover_table(doc: Document, program: Dict[str, Any], profile_pt: int) -> None:
    """Render cover info as invisible-border 2-column table.

    Accepts profile_pt (integer) to avoid depending on main() scope.
    """
    # build cover values from program
    date_text_original = program.get("date", "")
    date_text =format_date(date_text_original, "%Y年%m月%d日 %A", chinese_weekday=True, no_leading_zero=True)

    locations = program.get("locations") or []
    loc_text = ""
    if locations:
        loc_text = locations[0]
        if len(locations) > 1:
            loc_text += f"（{locations[1]}）"

    organizers_text = "、".join(program.get("organizers", [])) if program.get("organizers") else ""
    co_organizers_text = "、".join(program.get("coOrganizers", [])) if program.get("coOrganizers") else ""
    joint_organizers_text = "、".join(program.get("jointOrganizers", [])) if program.get("jointOrganizers") else ""

    instructors = program.get("instructors") or []
    instructor_text = ""
    if locations:
        instructor_text = instructors[0]
        if len(instructors) > 1:
            instructor_text += f"（{instructors[1]}）"

    # helper to remove borders
    def set_table_borders(table):
        tbl_pr = table._element.tblPr
        existing = tbl_pr.findall(qn("w:tblBorders"))
        for e in existing:
            tbl_pr.remove(e)
        borders = OxmlElement("w:tblBorders")
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = OxmlElement(f"w:{name}")
            node.set(qn("w:val"), "nil")
            borders.append(node)
        tbl_pr.append(borders)

    # create table
    cover_table = doc.add_table(rows=0, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(cover_table)




    # 關閉自動調整，轉為使用固定欄寬
    cover_table.autofit = False

    left_table_width_label = 3.5
    # 設定欄寬（範例：左欄 6 cm，右欄 10 cm）
    cover_table.columns[0].width = Cm(left_table_width_label)
    cover_table.columns[1].width = Cm(10)

    def add_cover_row(table, label, value):
        if value is None or (isinstance(value, str) and value.strip() == ""):
            return
        row = table.add_row().cells
        p_label = row[0].paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_label = p_label.add_run(label)
        set_run_font(run_label, profile_pt, bold=True)
        p_val = row[1].paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_val = p_val.add_run(value)
        set_run_font(run_val, profile_pt)

    # add rows (order)
    add_cover_row(cover_table, "日期：", date_text)
    add_cover_row(cover_table, "地點：", loc_text)
    add_cover_row(cover_table, "主辦單位：", organizers_text)
    add_cover_row(cover_table, "協辦單位：", co_organizers_text)
    add_cover_row(cover_table, "合辦單位：", joint_organizers_text)
    add_cover_row(cover_table, "指導單位：", instructor_text)


    # 如果 table 已有 row（或之後會新增 row），也對每個 cell 指定 width（保險做法）
    for row in cover_table.rows:
        row.cells[0].width = Cm(left_table_width_label)
        row.cells[1].width = Cm(10)
    # spacing after table
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Render program to docx")
    parser.add_argument("--program-id", type=int, default=None, help="Program id to render")
    parser.add_argument("--out", type=Path, default=None, help="Output .docx path")
    args = parser.parse_args()

    initialize()
    program = load_program(args.program_id)

    # Build chairs/speakers enriched with influencer data
    infl_file = DATA_DIR / "shared" / "influencer_data.json"
    try:
        influencers = json.loads(infl_file.read_text(encoding="utf-8"))
    except OSError:
        influencers = []
    chairs, speakers = build_people(program, influencers)

    schedule_rows = build_schedule(program)

    event_name = (program.get("eventNames") or ["Program"])[0]



    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal_font.size = Pt(12)

    # Heading style for TOC inclusion (set 置中)
    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    heading1.font.size = Pt(28)  # <- changed as requested
    heading1.font.bold = True
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    # Style constants (pt) derived from the HTML template
    TITLE_PT = 28
    NAME_PT = 18
    PROFILE_PT = 14
    TABLE_PT = 14  # <- changed as requested


    schedule_rows = build_schedule(program)
    if schedule_rows:


        for row in schedule_rows:


            doc = Document()
            name = row.get("speaker", "")
            title = row.get("topic", "")
            no = row.get("no", "")
            p    = doc.add_paragraph()
            name_run = p.add_run(name)
            set_run_font(name_run, TITLE_PT, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # title（同一行，普通字）
            doc.add_paragraph()
            p    = doc.add_paragraph()
            name_run = p.add_run(title)
            set_run_font(name_run, TITLE_PT, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            out_path =  (OUTPUT_DIR / f"{no}-1_{name}_{title}.docx")
            doc.save(out_path)
            print(f"Saved docx to {out_path}")









def update_docx_fields_with_word(docx_path: str, visible: bool = False) -> None:
    """
    Use MS Word COM to open the docx, update fields (TOC, page numbers), save and close.
    Requires Windows and MS Word installed. Install with: pip install pywin32
    """
    try:
        import pythoncom
        from win32com.client import Dispatch, constants
    except Exception as e:
        print("pywin32 未安裝或不可用：", e)
        return

    # Open Word
    word = Dispatch("Word.Application")
    word.Visible = visible  # True 可在執行時看到 Word 視窗（除錯用）
    # open document (Read/Write)
    doc = word.Documents.Open(str(docx_path))

    # 更新所有欄位（含 TOC）與所有目錄（若有複數 TOC）
    try:
        # Update fields generally
        doc.Fields.Update()
        # Update tables of contents specifically
        toc_count = doc.TablesOfContents.Count
        if toc_count > 0:
            for i in range(1, toc_count + 1):
                toc = doc.TablesOfContents(i)
                toc.Update()
    except Exception as e:
        print("更新欄位/TOC 時發生錯誤：", e)

    # Save and close
    doc.Save()
    doc.Close(False)
    word.Quit()


if __name__ == "__main__":
    main()
