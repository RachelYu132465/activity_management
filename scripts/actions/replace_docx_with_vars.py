# replace_docx_with_vars.py
# pip install python-docx
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import json
import argparse
import re
from typing import Any, Dict, Iterable, Tuple
from docx import Document

from scripts.core.bootstrap import TEMPLATE_DIR, DATA_DIR, OUTPUT_DIR

# ----- JSON path helpers -----
_path_token = re.compile(r"(?:\.?([^\.\[\]]+))(?:\[(\d+)\])?")

def get_by_path(obj: Any, path: str) -> Any:
    cur = obj
    for key, idx in _path_token.findall(path):
        if key:
            if not isinstance(cur, dict) or key not in cur:
                return None
            cur = cur[key]
        if idx:
            i = int(idx)
            if not isinstance(cur, list) or i >= len(cur):
                return None
            cur = cur[i]
    return cur

def iter_json_records(payload: Any):
    if isinstance(payload, list):
        for i, v in enumerate(payload):
            if isinstance(v, dict):
                yield ("[{}]".format(i), v)
    elif isinstance(payload, dict):
        yield ("", payload)
        for k, v in payload.items():
            if isinstance(v, list):
                for i, item in enumerate(v):
                    if isinstance(item, dict):
                        yield ("{}[{}]".format(k, i), item)

def iter_json_paths(node: Any, base: str = ""):
    if isinstance(node, dict):
        for k, v in node.items():
            key = "{}.{}".format(base, k) if base else k
            yield from iter_json_paths(v, key)
    elif isinstance(node, list):
        for i, v in enumerate(node):
            key = "{}[{}]".format(base, i)
            yield from iter_json_paths(v, key)
    else:
        yield (base, node)

def build_mapping_from_record(rec: dict) -> Dict[str, str]:
    m: Dict[str, str] = {}
    for path, val in iter_json_paths(rec):
        if isinstance(val, str):
            t = val.strip()
            if t:
                m.setdefault(t, "{{" + path + "}}")
    return m

# ----- DOCX replace -----
def replace_text_multi(text: str, mapping: Dict[str, str]) -> str:
    for target in sorted(mapping.keys(), key=len, reverse=True):
        text = text.replace(target, mapping[target])
    return text

def merge_runs(p):
    if not p.runs:
        return
    full = "".join(r.text for r in p.runs)
    for r in p.runs[1:]:
        p._element.remove(r._element)
    p.runs[0].text = full

def replace_in_paragraph(p, mapping: Dict[str, str], merge=False):
    changed = False
    for r in p.runs:
        new = replace_text_multi(r.text, mapping)
        if new != r.text:
            r.text = new
            changed = True
    if not changed and merge and p.runs:
        merge_runs(p)
        p.runs[0].text = replace_text_multi(p.runs[0].text, mapping)

def replace_in_document(doc: Document, mapping: Dict[str, str], merge_runs_if_needed=False):
    for para in doc.paragraphs:
        replace_in_paragraph(para, mapping, merge_runs_if_needed)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, mapping, merge_runs_if_needed)

# ----- main -----
def main():
    ap = argparse.ArgumentParser(
        description="指定 JSON 檔名，於 Data/** 遞迴搜尋所有同名 JSON，合併值→{{路徑}} 對照，替換 DOCX。"
    )
    ap.add_argument("--template", required=True, help="DOCX 檔名（例如 notice.docx），於 Templates/** 遞迴尋找")
    ap.add_argument("--json-file", required=True, help="要匹配的 JSON 檔名（例如 program_data.json），於 Data/** 遞迴尋找多個")
    ap.add_argument("--id", type=int, help="若 JSON 是清單，選擇 id 等於此值的那筆")
    ap.add_argument("--filter-path", help="JSON 路徑（如 eventNames、eventNames[0]、planName）")
    ap.add_argument("--filter-equals", help="值必須等於此字串")
    ap.add_argument("--filter-contains", help="值需包含此子字串（支援 list/str）")
    ap.add_argument("--merge-runs", action="store_true", help="需要時合併段落 runs 再替換（可能失去局部粗斜體/顏色）")
    args = ap.parse_args()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # 找模板
    t_matches = list(TEMPLATE_DIR.rglob(args.template))
    if not t_matches:
        raise SystemExit("找不到模板：{} 於 {} 下".format(args.template, TEMPLATE_DIR))
    input_docx = t_matches[0]

    # 找所有同名 JSON
    j_matches = list(DATA_DIR.rglob(args.json_file))
    if not j_matches:
        raise SystemExit("找不到 JSON：{} 於 {} 下".format(args.json_file, DATA_DIR))

    # 由多個 JSON 檔合併 mapping
    mapping: Dict[str, str] = {}
    picked_sources = []
    for jf in j_matches:
        try:
            payload = json.loads(jf.read_text(encoding="utf-8"))
        except Exception as e:
            print("[WARN] 讀取失敗 {}: {}".format(jf, e))
            continue

        # 依使用者條件在檔內挑選一筆（或多筆）record
        def record_ok(rec: dict) -> bool:
            if args.id is not None and rec.get("id") != args.id:
                return False
            if args.filter_path and (args.filter_equals is not None or args.filter_contains is not None):
                val = get_by_path(rec, args.filter_path)
                if args.filter_equals is not None:
                    return val == args.filter_equals
                if args.filter_contains is not None:
                    if isinstance(val, str):
                        return args.filter_contains in val
                    if isinstance(val, list):
                        return any(isinstance(x, str) and args.filter_contains in x for x in val)
                    return False
            return True

        picked_any = False
        for root_path, rec in iter_json_records(payload):
            if isinstance(rec, dict) and record_ok(rec):
                m = build_mapping_from_record(rec)
                # 合併：保留先到者（避免覆蓋）；若要覆蓋，改成 mapping.update(m)
                for k, v in m.items():
                    mapping.setdefault(k, v)
                picked_any = True

        if picked_any:
            picked_sources.append(str(jf))

    if not mapping:
        raise SystemExit("沒有找到符合條件的 JSON 值可替換（請檢查 --id / --filter-* 與 json-file 是否正確）。")

    print("[INFO] 使用到的 JSON：")
    for s in picked_sources:
        print("  -", s)
    print("[INFO] 產生對照筆數：", len(mapping))

    # 套用到 DOCX
    doc = Document(input_docx)
    replace_in_document(doc, mapping, merge_runs_if_needed=args.merge_runs)
    out_path = OUTPUT_DIR / (input_docx.stem + "_template.docx")
    doc.save(out_path)
    print("[OK] 已輸出：{}".format(out_path))

if __name__ == "__main__":
    main()
